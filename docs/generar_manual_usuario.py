# -*- coding: utf-8 -*-
"""
Generador del Manual de Usuario - StoreHub
Produce MANUAL_USUARIO_StoreHub.docx, completamente explicativo para
usuarios finales (clientes y administradores) sin conocimientos técnicos previos.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)
    run.font.size = Pt(18)
    run.bold = True

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.font.size = Pt(14)
    run.bold = True

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x79)
    run.font.size = Pt(12)
    run.bold = True

def body(doc, text, bold=False, italic=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        rb.bold = True
        rb.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def number_item(doc, num, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.35)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        rb.bold = True
        rb.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def tip_box(doc, text, kind="tip"):
    """Crea un recuadro de consejo/advertencia/nota."""
    colors = {
        "tip":     ("E8F5E9", "2E7D32", "CONSEJO"),
        "warning": ("FFF8E1", "F57F17", "IMPORTANTE"),
        "error":   ("FFEBEE", "C62828", "ADVERTENCIA"),
        "info":    ("E3F2FD", "1565C0", "NOTA"),
    }
    bg, fg_hex, label = colors.get(kind, colors["info"])
    fg = tuple(int(fg_hex[i:i+2], 16) for i in (0, 2, 4))
    p = doc.add_paragraph()
    r_label = p.add_run(f"  {label}:  ")
    r_label.bold = True
    r_label.font.size = Pt(10)
    r_label.font.color.rgb = RGBColor(*fg)
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def sep(doc):
    doc.add_paragraph()

def tbl_header(doc, cols, widths=None):
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, col in enumerate(cols):
        cell = hdr.cells[i]
        cell.text = col
        set_cell_bg(cell, "1A569C")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for i, w in enumerate(widths):
            tbl.columns[i].width = Cm(w)
    return tbl

def add_row(tbl, vals, bg=None):
    row = tbl.add_row()
    for i, v in enumerate(vals):
        cell = row.cells[i]
        cell.text = str(v)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        if bg:
            set_cell_bg(cell, bg)
    return row

# ─── build ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ═══════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════
    sep(doc); sep(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MANUAL DE USUARIO")
    r.bold = True; r.font.size = Pt(32)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("StoreHub — Plataforma de Compras en Línea")
    r2.bold = True; r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    sep(doc)

    for lbl, val in [
        ("Versión:", "1.0.0"),
        ("Fecha:", "04 de Marzo de 2026"),
        ("Dirigido a:", "Clientes y Administradores"),
        ("Plataforma:", "https://inventory-2-sewi.onrender.com"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{lbl}  "); r1.bold = True; r1.font.size = Pt(12)
        r2 = p.add_run(val); r2.font.size = Pt(12)
    sep(doc)
    p_a = doc.add_paragraph(); p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = p_a.add_run("Autores: "); ra.bold = True; ra.font.size = Pt(11)
    p_a.add_run("Sebastián Gallego V.   Santiago Sánchez S.   Duván Ochoa H.   Ricky Lotero S."
                ).font.size = Pt(11)
    p_c = doc.add_paragraph(); p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_c.add_run("Cliente: "); rc.bold = True; rc.font.size = Pt(11)
    p_c.add_run("Julián Alberto Garzón García / Distribuciones S y J").font.size = Pt(11)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # TABLA DE CONTENIDOS
    # ═══════════════════════════════════════════
    h1(doc, "Tabla de Contenidos")
    toc = [
        "SECCIÓN 1 — Introducción y Requisitos Previos",
        "  1.1  ¿Qué es StoreHub?",
        "  1.2  ¿A quién va dirigido este manual?",
        "  1.3  Requisitos de Hardware",
        "  1.4  Requisitos de Software (Navegador Web)",
        "  1.5  Conexión a Internet",
        "  1.6  Acceso a la Plataforma",
        "SECCIÓN 2 — Registro e Inicio de Sesión (Clientes)",
        "  2.1  Crear una cuenta nueva",
        "  2.2  Iniciar sesión",
        "  2.3  Recuperar o cambiar contraseña",
        "  2.4  Cerrar sesión",
        "SECCIÓN 3 — Navegación por la Tienda",
        "  3.1  Página principal",
        "  3.2  Explorar categorías",
        "  3.3  Buscar productos",
        "  3.4  Ver detalle de un producto",
        "  3.5  Sección de Ofertas",
        "SECCIÓN 4 — Carrito de Compras",
        "  4.1  Agregar un producto al carrito",
        "  4.2  Ver el carrito",
        "  4.3  Modificar cantidades",
        "  4.4  Eliminar productos del carrito",
        "  4.5  Vaciar el carrito",
        "SECCIÓN 5 — Proceso de Compra (Checkout)",
        "  5.1  Gestionar Direcciones de Envío",
        "  5.2  Confirmar la orden",
        "  5.3  Pantalla de confirmación",
        "SECCIÓN 6 — Seguimiento de Pedidos",
        "  6.1  Ver mis pedidos",
        "  6.2  Estados del pedido",
        "  6.3  Detalle de un pedido",
        "SECCIÓN 7 — Perfil de Usuario",
        "  7.1  Ver y editar mis datos",
        "SECCIÓN 8 — Panel de Administración (Solo ADMIN)",
        "  8.1  Acceso al panel",
        "  8.2  Dashboard",
        "  8.3  Gestión de Pedidos",
        "  8.4  Gestión de Productos",
        "  8.5  Control de Inventario",
        "  8.6  Categorías",
        "  8.7  Usuarios",
        "  8.8  Ofertas",
        "  8.9  Reportes",
        "SECCIÓN 9 — Preguntas Frecuentes y Solución de Problemas",
        "SECCIÓN 10 — Glosario de Términos",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 1 — INTRODUCCIÓN Y REQUISITOS
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 1 — Introducción y Requisitos Previos")

    h2(doc, "1.1 ¿Qué es StoreHub?")
    body(doc,
        "StoreHub es una plataforma de comercio electrónico (tienda en línea) que le permite "
        "comprar productos desde la comodidad de su hogar, utilizando únicamente un navegador "
        "de internet. No es necesario instalar ningún programa en su computadora o celular."
    )
    body(doc,
        "Si usted es administrador de la tienda, StoreHub también le proporciona un panel "
        "completo para gestionar los productos, pedidos, inventario, usuarios y reportes de ventas."
    )
    sep(doc)
    body(doc, "En resumen, StoreHub tiene dos modos de uso:", bold=True)
    tbl_modos = tbl_header(doc, ["Tipo de usuario", "Qué puede hacer"])
    add_row(tbl_modos, ["Cliente",
        "Explorar productos, agregar al carrito, realizar compras, rastrear pedidos."])
    add_row(tbl_modos, ["Administrador",
        "Gestionar productos, categorías, pedidos, inventario, usuarios, ofertas y reportes."])
    sep(doc)

    h2(doc, "1.2 ¿A quién va dirigido este manual?")
    body(doc,
        "Este manual está escrito para cualquier persona que use StoreHub, "
        "independientemente de sus conocimientos tecnológicos. No se requiere experiencia "
        "previa con plataformas de compras en línea. Cada paso está explicado de forma "
        "clara y detallada."
    )
    bullet(doc, "Clientes: personas que desean comprar productos en la tienda.")
    bullet(doc, "Administradores: personal encargado de operar y gestionar la tienda.")
    sep(doc)

    h2(doc, "1.3 Requisitos de Hardware")
    body(doc,
        "Para acceder a StoreHub no necesita un equipo de alta gama. Cualquier dispositivo "
        "moderno con acceso a internet es suficiente. A continuación se detallan los "
        "requisitos mínimos y recomendados:"
    )
    sep(doc)
    body(doc, "Computadora de Escritorio o Portátil (PC / Mac / Linux)", bold=True)
    tbl_pc = tbl_header(doc, ["Componente", "Mínimo", "Recomendado"])
    for row in [
        ("Procesador",    "Intel Pentium / AMD A-Series (1.6 GHz)",    "Intel Core i3 / AMD Ryzen 3 o superior"),
        ("Memoria RAM",   "2 GB",                                       "4 GB o más"),
        ("Pantalla",      "Resolución 1024 × 768 (10 pulgadas)",        "1280 × 720 o superior"),
        ("Almacenamiento","No se requiere espacio local",               "No aplica (todo es en línea)"),
        ("Conexión",      "Wi-Fi o cable Ethernet",                     "Cable Ethernet o Wi-Fi 5 GHz"),
    ]:
        add_row(tbl_pc, row)
    sep(doc)

    body(doc, "Dispositivo Móvil (Teléfono / Tableta)", bold=True)
    tbl_mob = tbl_header(doc, ["Componente", "Mínimo", "Recomendado"])
    for row in [
        ("Sistema Operativo", "Android 8.0 / iOS 13",              "Android 11+ / iOS 15+"),
        ("Pantalla",          "4.5 pulgadas",                       "5.5 pulgadas o más"),
        ("RAM",               "2 GB",                               "3 GB o más"),
        ("Conexión",          "4G / Wi-Fi",                         "4G LTE / Wi-Fi"),
    ]:
        add_row(tbl_mob, row)
    sep(doc)
    tip_box(doc,
        "StoreHub funciona correctamente en dispositivos de hace 4-5 años. "
        "Si puede abrir páginas como Google o YouTube, puede usar StoreHub sin problemas.",
        "tip"
    )
    sep(doc)

    h2(doc, "1.4 Requisitos de Software (Navegador Web)")
    body(doc,
        "No necesita instalar ningún programa especial. Solo necesita un navegador web "
        "(el programa que usa para navegar en internet). Los siguientes navegadores son "
        "compatibles con StoreHub:"
    )
    tbl_nav = tbl_header(doc, ["Navegador", "Versión Mínima", "¿Gratuito?", "Descarga"])
    for row in [
        ("Google Chrome",  "90 o superior", "Sí", "google.com/chrome"),
        ("Mozilla Firefox","88 o superior", "Sí", "mozilla.org/firefox"),
        ("Microsoft Edge", "91 o superior", "Sí", "microsoft.com/edge"),
        ("Safari (Mac/iOS)","14 o superior","Sí", "Incluido en Mac/iPhone"),
        ("Opera",          "76 o superior", "Sí", "opera.com"),
    ]:
        add_row(tbl_nav, row)
    sep(doc)
    tip_box(doc,
        "Se recomienda usar Google Chrome o Microsoft Edge para la mejor experiencia. "
        "Asegúrese de que su navegador esté actualizado: en Chrome vaya a los tres puntos "
        "(arriba a la derecha) → Ayuda → Información de Google Chrome.",
        "tip"
    )
    tip_box(doc,
        "NO use Internet Explorer (IE). Este navegador es obsoleto y puede hacer que la "
        "página no funcione correctamente.",
        "warning"
    )
    sep(doc)

    h2(doc, "1.5 Conexión a Internet")
    body(doc,
        "StoreHub es una aplicación completamente en línea. Necesita estar conectado a "
        "internet para usarla. A continuación se explica qué velocidad de conexión necesita:"
    )
    tbl_net = tbl_header(doc, ["Actividad", "Velocidad Mínima Requerida"])
    for row in [
        ("Ver productos y navegar la tienda",   "1 Mbps de bajada"),
        ("Cargar imágenes de productos",         "2 Mbps de bajada"),
        ("Realizar un pedido",                   "1 Mbps de bajada"),
        ("Panel de administración (reportes)",   "2 Mbps de bajada"),
    ]:
        add_row(tbl_net, row)
    sep(doc)
    tip_box(doc,
        "La velocidad promedio de internet en Colombia es de 50–100 Mbps, "
        "lo que es más que suficiente para usar StoreHub sin inconvenientes.",
        "info"
    )
    tip_box(doc,
        "Si nota que la página carga lento, verifique su conexión a internet antes de "
        "reportar un problema. Puede probar su velocidad en speedtest.net.",
        "tip"
    )
    sep(doc)

    h2(doc, "1.6 Acceso a la Plataforma")
    body(doc,
        "Para acceder a StoreHub, abra su navegador web y escriba la siguiente dirección "
        "en la barra de direcciones (la parte de arriba del navegador donde dice 'https://'):"
    )
    sep(doc)
    p_url = doc.add_paragraph()
    p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_url = p_url.add_run("https://inventory-2-sewi.onrender.com")
    r_url.bold = True; r_url.font.size = Pt(14)
    r_url.font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)
    sep(doc)
    tip_box(doc,
        "Guarde esta dirección en sus favoritos o marcadores para acceder más rápido "
        "la próxima vez. En Chrome: presione Ctrl+D (Windows) o Cmd+D (Mac).",
        "tip"
    )
    sep(doc)
    body(doc, "Al ingresar a la plataforma verá la página principal de la tienda, que muestra:", bold=True)
    bullet(doc, "El logotipo y nombre de la tienda (StoreHub) en la parte superior.")
    bullet(doc, "Una barra de navegación con categorías de productos.")
    bullet(doc, "Un banner o imagen principal (Hero) con promociones destacadas.")
    bullet(doc, "La cuadrícula de productos disponibles.")
    bullet(doc, "El ícono del carrito de compras (esquina superior derecha).")
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 2 — REGISTRO E INICIO DE SESIÓN
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 2 — Registro e Inicio de Sesión (Clientes)")
    body(doc,
        "Para comprar en StoreHub necesita crear una cuenta gratuita. Si ya tiene una cuenta, "
        "puede iniciar sesión directamente. Si solo quiere explorar los productos, puede hacerlo "
        "sin registrarse."
    )
    sep(doc)

    h2(doc, "2.1 Crear una cuenta nueva")
    body(doc,
        "Crear una cuenta en StoreHub es completamente gratuito y solo toma unos pocos minutos."
    )
    body(doc, "Pasos para registrarse:", bold=True)
    number_item(doc, 1, "Abra StoreHub en su navegador (ver dirección en sección 1.6).")
    number_item(doc, 2, 'Haga clic en el botón "Iniciar Sesión" ubicado en la esquina superior derecha de la página.')
    number_item(doc, 3, 'En la página de inicio de sesión, busque el enlace "¿No tienes cuenta? Regístrate" y haga clic en él.')
    number_item(doc, 4, "Complete el formulario de registro con los siguientes datos:")
    bullet(doc, "Nombre completo: su nombre y apellido.", level=1)
    bullet(doc, "Correo electrónico: una dirección de email válida (ej: nombre@gmail.com). Esta será su usuario.", level=1)
    bullet(doc, "Contraseña: elija una contraseña de mínimo 6 caracteres. Use letras y números para mayor seguridad.", level=1)
    bullet(doc, "Teléfono (opcional): su número de contacto.", level=1)
    number_item(doc, 5, 'Haga clic en el botón "Registrarse" o "Crear cuenta".')
    number_item(doc, 6, "Si todos los datos son correctos, verá un mensaje de bienvenida y será redirigido automáticamente a la tienda.")
    sep(doc)
    tip_box(doc,
        "Use un correo electrónico real al que tenga acceso, ya que lo necesitará para "
        "iniciar sesión y para recuperar su cuenta si olvida la contraseña.",
        "warning"
    )
    tip_box(doc,
        "El correo electrónico no puede estar registrado previamente. Si ve el mensaje "
        '"Este correo ya está registrado", significa que ya tiene una cuenta. '
        "Intente iniciar sesión con ese correo.",
        "info"
    )
    sep(doc)

    h2(doc, "2.2 Iniciar sesión")
    body(doc,
        "Si ya tiene una cuenta, siga estos pasos para ingresar a StoreHub:"
    )
    number_item(doc, 1, 'Haga clic en "Iniciar Sesión" en la esquina superior derecha.')
    number_item(doc, 2, "Escriba su correo electrónico registrado en el campo 'Email'.")
    number_item(doc, 3, "Escriba su contraseña en el campo 'Contraseña'.")
    number_item(doc, 4, 'Haga clic en el botón "Ingresar".')
    number_item(doc, 5, "Si los datos son correctos, verá un mensaje de bienvenida y será redirigido a la tienda (si es cliente) o al panel de administración (si es administrador).")
    sep(doc)
    tip_box(doc,
        "Su sesión se mantiene activa durante 1 hora sin actividad. Después de ese tiempo, "
        "el sistema le pedirá que inicie sesión nuevamente por razones de seguridad.",
        "info"
    )
    tip_box(doc,
        "Si ve el mensaje 'Credenciales incorrectas', verifique que su correo y contraseña "
        "sean correctos. Recuerde que la contraseña distingue entre mayúsculas y minúsculas.",
        "warning"
    )
    sep(doc)

    h2(doc, "2.3 Recuperar o cambiar contraseña")
    body(doc,
        "Si olvidó su contraseña, puede actualizarla desde su perfil de usuario una vez que "
        "inicie sesión. Si no puede acceder a su cuenta, contacte al administrador de la tienda."
    )
    body(doc, "Para cambiar su contraseña (estando dentro de su cuenta):", bold=True)
    number_item(doc, 1, "Inicie sesión normalmente.")
    number_item(doc, 2, "Haga clic en su nombre de usuario (esquina superior derecha).")
    number_item(doc, 3, 'Seleccione "Mi Perfil" o "Configuración".')
    number_item(doc, 4, "En el campo 'Contraseña', escriba la nueva contraseña que desea usar.")
    number_item(doc, 5, 'Haga clic en "Guardar cambios".')
    sep(doc)

    h2(doc, "2.4 Cerrar sesión")
    body(doc, "Para cerrar sesión de forma segura:")
    number_item(doc, 1, "Haga clic en su nombre o en el ícono de usuario (parte superior derecha).")
    number_item(doc, 2, 'Haga clic en "Cerrar sesión" o "Salir".')
    number_item(doc, 3, "Verá un mensaje de confirmación. Será redirigido a la página de inicio de sesión.")
    sep(doc)
    tip_box(doc,
        "Siempre cierre sesión cuando use un computador o dispositivo compartido "
        "(biblioteca, cibercafé, trabajo). Esto protege su cuenta y sus datos personales.",
        "warning"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 3 — NAVEGACIÓN
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 3 — Navegación por la Tienda")

    h2(doc, "3.1 Página principal")
    body(doc,
        "Cuando abre StoreHub, lo primero que ve es la página principal. Esta es su punto "
        "de partida para explorar todos los productos disponibles."
    )
    body(doc, "La página principal está compuesta por:", bold=True)
    bullet(doc, "Barra superior (Header): contiene el logotipo de StoreHub, las categorías de productos, "
                "el botón de iniciar sesión y el ícono del carrito de compras.")
    bullet(doc, "Banner principal (Hero): imagen o mensaje de bienvenida con las promociones más destacadas.")
    bullet(doc, "Filtro de categorías: botones para filtrar los productos por categoría (ej: Electrónica, Ropa, etc.).")
    bullet(doc, "Cuadrícula de productos: muestra las tarjetas de todos los productos disponibles.")
    bullet(doc, "Pie de página (Footer): información de contacto, términos y condiciones, política de privacidad.")
    sep(doc)
    tip_box(doc,
        "Si accede desde un celular, la barra de categorías se convierte en un menú "
        "desplegable. Toque el ícono de tres líneas (≡) para abrirlo.",
        "info"
    )
    sep(doc)

    h2(doc, "3.2 Explorar categorías")
    body(doc,
        "Los productos están organizados en categorías para que sea más fácil encontrar "
        "lo que busca. Por ejemplo: 'Cafés', 'Snacks', 'Bebidas', etc."
    )
    body(doc, "¿Cómo filtrar por categoría?", bold=True)
    number_item(doc, 1, "En la barra de navegación superior o en el filtro de categorías, verá botones con los nombres de las categorías disponibles.")
    number_item(doc, 2, "Haga clic en la categoría que le interesa.")
    number_item(doc, 3, "La cuadrícula de productos se actualizará automáticamente mostrando solo los productos de esa categoría.")
    number_item(doc, 4, 'Para ver todos los productos de nuevo, haga clic en "Todos" o "Ver todo".')
    sep(doc)

    h2(doc, "3.3 Buscar productos")
    body(doc,
        "Puede buscar productos específicos escribiendo el nombre en la barra de búsqueda "
        "ubicada en la parte superior de la página."
    )
    body(doc, "Pasos para buscar un producto:", bold=True)
    number_item(doc, 1, "Ubique la barra de búsqueda (generalmente tiene un ícono de lupa) en la parte superior.")
    number_item(doc, 2, "Escriba el nombre del producto o una palabra clave (ej: 'café', 'chocolate', 'crema').")
    number_item(doc, 3, "Los resultados se mostrarán automáticamente mientras escribe, o presione Enter para buscar.")
    number_item(doc, 4, "Si no se encuentran resultados, intente con palabras más generales.")
    sep(doc)
    tip_box(doc,
        "La búsqueda no distingue entre mayúsculas y minúsculas. 'CAFÉ' y 'café' "
        "darán los mismos resultados.",
        "tip"
    )
    sep(doc)

    h2(doc, "3.4 Ver detalle de un producto")
    body(doc,
        "Al hacer clic sobre cualquier tarjeta de producto, accederá a la página de "
        "detalle donde encontrará toda la información del artículo."
    )
    body(doc, "En la página de detalle encontrará:", bold=True)
    bullet(doc, "Imagen del producto (puede haber varias fotos).")
    bullet(doc, "Nombre completo del producto.")
    bullet(doc, "Descripción detallada: características, peso, ingredientes, etc.")
    bullet(doc, "Precio unitario del producto.")
    bullet(doc, "Disponibilidad (cuántas unidades hay en stock).")
    bullet(doc, "Selector de cantidad: indique cuántas unidades desea comprar.")
    bullet(doc, 'Botón "Agregar al carrito": para añadir el producto a su carrito de compras.')
    bullet(doc, "Categoría a la que pertenece el producto.")
    sep(doc)
    tip_box(doc,
        "Si el stock de un producto muestra 0 o aparece como 'Agotado', "
        "no podrá agregarlo al carrito hasta que el administrador reponga el inventario.",
        "warning"
    )
    sep(doc)

    h2(doc, "3.5 Sección de Ofertas")
    body(doc,
        "StoreHub cuenta con una sección exclusiva de Ofertas donde encontrará "
        "productos con precios especiales o promociones temporales."
    )
    number_item(doc, 1, 'Haga clic en "Ofertas" en la barra de navegación superior.')
    number_item(doc, 2, "Verá la lista de productos con descuento activo.")
    number_item(doc, 3, "El precio con descuento y el porcentaje de ahorro se muestran claramente.")
    number_item(doc, 4, "Para aprovechar una oferta, agregue el producto al carrito normalmente.")
    sep(doc)
    tip_box(doc,
        "Las ofertas son temporales. Un producto puede dejar de estar en oferta "
        "en cualquier momento, así que no espere demasiado si le interesa.",
        "info"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 4 — CARRITO
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 4 — Carrito de Compras")
    body(doc,
        "El carrito de compras es el espacio donde se guardan los productos que desea "
        "comprar antes de finalizar el pedido. Funciona igual que un carrito físico en "
        "un supermercado: puede agregar, quitar o modificar los productos antes de pagar."
    )
    sep(doc)
    tip_box(doc,
        "Debe haber iniciado sesión para poder agregar productos al carrito y "
        "realizar compras. Si no lo ha hecho, el sistema le pedirá que inicie sesión.",
        "warning"
    )
    sep(doc)

    h2(doc, "4.1 Agregar un producto al carrito")
    body(doc, "Hay dos formas de agregar un producto al carrito:")
    body(doc, "Opción A — Desde la cuadrícula de productos (página principal):", bold=True)
    number_item(doc, 1, "Encuentre el producto que desea en la página principal o en los resultados de búsqueda.")
    number_item(doc, 2, 'Haga clic en el botón "Agregar al carrito" que aparece en la tarjeta del producto.')
    number_item(doc, 3, "Verá una notificación indicando que el producto fue agregado correctamente.")
    sep(doc)
    body(doc, "Opción B — Desde la página de detalle del producto:", bold=True)
    number_item(doc, 1, "Haga clic en la tarjeta del producto para abrir su página de detalle.")
    number_item(doc, 2, "Use el selector de cantidad (+/-) para elegir cuántas unidades quiere.")
    number_item(doc, 3, 'Haga clic en "Agregar al carrito".')
    sep(doc)

    h2(doc, "4.2 Ver el carrito")
    number_item(doc, 1, "Haga clic en el ícono del carrito (parte superior derecha de la pantalla). Tiene forma de bolsa o carrito de compras.")
    number_item(doc, 2, "Se abrirá un panel lateral (o página) mostrando todos los productos que ha agregado.")
    number_item(doc, 3, "Verá el nombre de cada producto, la cantidad, el precio unitario y el subtotal.")
    number_item(doc, 4, "En la parte inferior verá el TOTAL de su compra.")
    sep(doc)
    tip_box(doc,
        "El número que aparece sobre el ícono del carrito indica cuántos artículos "
        "distintos tiene en él.",
        "info"
    )
    sep(doc)

    h2(doc, "4.3 Modificar cantidades")
    body(doc, "Para cambiar la cantidad de un producto ya agregado al carrito:")
    number_item(doc, 1, "Abra el carrito haciendo clic en el ícono.")
    number_item(doc, 2, 'Junto a cada producto verá botones "+" y "–".')
    number_item(doc, 3, 'Haga clic en "+" para aumentar la cantidad o "–" para disminuirla.')
    number_item(doc, 4, "El total se actualizará automáticamente.")
    sep(doc)
    tip_box(doc,
        "No puede agregar más unidades de las que hay en stock. Si intenta superar "
        "el límite disponible, el sistema le mostrará un aviso.",
        "warning"
    )
    sep(doc)

    h2(doc, "4.4 Eliminar productos del carrito")
    body(doc, "Para quitar un producto del carrito:")
    number_item(doc, 1, "Abra el carrito.")
    number_item(doc, 2, 'Haga clic en el ícono de basura (🗑) o en el botón "Eliminar" que aparece junto al producto.')
    number_item(doc, 3, "El producto será removido y el total se recalculará.")
    sep(doc)

    h2(doc, "4.5 Vaciar el carrito")
    body(doc, "Si desea empezar desde cero y eliminar todos los productos a la vez:")
    number_item(doc, 1, "Abra el carrito.")
    number_item(doc, 2, 'Busque y haga clic en el botón "Vaciar carrito".')
    number_item(doc, 3, "Se le pedirá confirmación. Acepte para eliminar todos los productos.")
    sep(doc)
    tip_box(doc,
        "Su carrito se guarda en la base de datos, por lo que si cierra el navegador "
        "y vuelve a iniciar sesión, sus productos seguirán allí.",
        "tip"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 5 — CHECKOUT
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 5 — Proceso de Compra (Checkout)")
    body(doc,
        "El proceso de compra (también llamado 'checkout') es el conjunto de pasos "
        "que debe seguir para convertir los productos de su carrito en un pedido real. "
        "Es muy sencillo y toma menos de 2 minutos."
    )
    sep(doc)

    h2(doc, "5.1 Gestionar Direcciones de Envío")
    body(doc,
        "Antes de confirmar su compra, necesita tener al menos una dirección de envío "
        "registrada. Esta es la dirección a donde llegará su pedido."
    )
    body(doc, "¿Cómo agregar una dirección?", bold=True)
    number_item(doc, 1, "Una vez que inicie sesión, diríjase al proceso de checkout (botón 'Finalizar compra' en el carrito).")
    number_item(doc, 2, 'Haga clic en "Agregar nueva dirección".')
    number_item(doc, 3, "Complete el formulario:")
    bullet(doc, "Calle / Carrera / Dirección: la dirección exacta (ej: Calle 45 # 23-10).", level=1)
    bullet(doc, "Ciudad: la ciudad donde vive (ej: Medellín).", level=1)
    bullet(doc, "Departamento: el departamento correspondiente (ej: Antioquia).", level=1)
    bullet(doc, "Información adicional (opcional): apartamento, piso, punto de referencia.", level=1)
    number_item(doc, 4, 'Haga clic en "Guardar dirección".')
    sep(doc)
    tip_box(doc,
        "Puede guardar varias direcciones y elegir en cada compra cuál usar. "
        "Por ejemplo, puede tener la dirección de su casa y la de su trabajo.",
        "tip"
    )
    sep(doc)

    h2(doc, "5.2 Confirmar la orden")
    body(doc, "Una vez que tiene su carrito listo y una dirección registrada:")
    number_item(doc, 1, 'Abra el carrito y haga clic en "Finalizar compra" o "Proceder al pago".')
    number_item(doc, 2, "Revise el resumen de su pedido: productos, cantidades y precios.")
    number_item(doc, 3, "Seleccione la dirección de envío donde desea recibir el pedido.")
    number_item(doc, 4, "Verifique el total final de su compra.")
    number_item(doc, 5, 'Haga clic en "Confirmar Pedido".')
    sep(doc)
    tip_box(doc,
        "Revise cuidadosamente los productos y la dirección antes de confirmar. "
        "Una vez confirmado el pedido, los cambios deben ser gestionados por el administrador.",
        "warning"
    )
    sep(doc)

    h2(doc, "5.3 Pantalla de confirmación")
    body(doc,
        "Después de confirmar el pedido, aparecerá una pantalla de confirmación que muestra:"
    )
    bullet(doc, "El número único de su pedido (guárdelo para hacer seguimiento).")
    bullet(doc, "El resumen de los productos comprados.")
    bullet(doc, "La dirección de envío seleccionada.")
    bullet(doc, "El total pagado.")
    bullet(doc, "El estado inicial del pedido: PENDIENTE.")
    sep(doc)
    tip_box(doc,
        "Le recomendamos tomar una foto o captura de pantalla de la confirmación "
        "como comprobante de su pedido.",
        "tip"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 6 — SEGUIMIENTO DE PEDIDOS
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 6 — Seguimiento de Pedidos")
    body(doc,
        "Puede hacer seguimiento a todos sus pedidos en cualquier momento desde su cuenta. "
        "Sabrá exactamente en qué estado está cada uno."
    )
    sep(doc)

    h2(doc, "6.1 Ver mis pedidos")
    number_item(doc, 1, "Inicie sesión en su cuenta.")
    number_item(doc, 2, 'Haga clic en su nombre de usuario (esquina superior derecha) y seleccione "Mis Pedidos", o busque el enlace "Pedidos" en la barra de navegación.')
    number_item(doc, 3, "Verá la lista de todos sus pedidos ordenados del más reciente al más antiguo.")
    number_item(doc, 4, "Cada pedido muestra: número de pedido, fecha, total, estado actual y los productos incluidos.")
    sep(doc)

    h2(doc, "6.2 Estados del pedido")
    body(doc, "Cada pedido pasa por diferentes estados a medida que es procesado:")
    tbl_estados = tbl_header(doc, ["Estado", "Significado", "¿Qué esperar?"])
    for row in [
        ("PENDIENTE",
         "El pedido fue recibido y está esperando ser procesado.",
         "El administrador revisará su pedido pronto."),
        ("EN PREPARACIÓN",
         "El equipo está preparando su pedido (empacando, verificando stock).",
         "Su pedido está siendo alistado para envío."),
        ("ENTREGADO",
         "El pedido fue entregado exitosamente.",
         "Confirme la recepción de todos los productos."),
        ("CANCELADO",
         "El pedido fue cancelado (por el cliente o el administrador).",
         "Comuníquese con la tienda si tiene dudas."),
    ]:
        add_row(tbl_estados, row)
    sep(doc)

    h2(doc, "6.3 Detalle de un pedido")
    number_item(doc, 1, "En la lista de pedidos, haga clic sobre el número de pedido o en 'Ver detalle'.")
    number_item(doc, 2, "Verá la información completa:")
    bullet(doc, "Número de pedido.", level=1)
    bullet(doc, "Fecha y hora en que se realizó.", level=1)
    bullet(doc, "Lista de productos con nombre, cantidad y precio unitario.", level=1)
    bullet(doc, "Subtotal por producto y total general.", level=1)
    bullet(doc, "Dirección de envío.", level=1)
    bullet(doc, "Estado actual del pedido.", level=1)
    sep(doc)
    tip_box(doc,
        "Si tiene algún problema con su pedido (producto incorrecto, cantidad errónea), "
        "comuníquese con el equipo de la tienda indicando su número de pedido.",
        "info"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 7 — PERFIL
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 7 — Perfil de Usuario")

    h2(doc, "7.1 Ver y editar mis datos")
    body(doc,
        "Puede actualizar su información personal en cualquier momento desde su perfil."
    )
    number_item(doc, 1, "Inicie sesión.")
    number_item(doc, 2, "Haga clic en su nombre (esquina superior derecha).")
    number_item(doc, 3, 'Seleccione "Mi Perfil" o "Editar perfil".')
    number_item(doc, 4, "Podrá modificar los siguientes campos:")
    bullet(doc, "Nombre completo.", level=1)
    bullet(doc, "Número de teléfono.", level=1)
    bullet(doc, "Contraseña (escriba la nueva contraseña en el campo correspondiente).", level=1)
    number_item(doc, 5, 'Haga clic en "Guardar cambios" para aplicar las modificaciones.')
    sep(doc)
    tip_box(doc,
        "El correo electrónico no se puede cambiar una vez registrado, ya que es "
        "el identificador único de su cuenta.",
        "info"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 8 — PANEL DE ADMINISTRACIÓN
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 8 — Panel de Administración (Solo ADMIN)")
    body(doc,
        "Esta sección está destinada exclusivamente para los usuarios con rol de Administrador. "
        "Si usted es cliente, puede saltar a la Sección 9."
    )
    sep(doc)
    tip_box(doc,
        "Si intenta acceder al panel de administración sin tener permisos de administrador, "
        "el sistema le mostrará un mensaje de 'Acceso denegado' y le redirigirá al inicio.",
        "error"
    )
    sep(doc)

    h2(doc, "8.1 Acceso al panel de administración")
    body(doc,
        "El panel de administración es una área exclusiva donde puede gestionar "
        "todos los aspectos de la tienda."
    )
    number_item(doc, 1, "Inicie sesión con su cuenta de administrador (correo y contraseña de admin).")
    number_item(doc, 2, "El sistema detectará automáticamente su rol y lo redirigirá al panel de administración.")
    number_item(doc, 3, "También puede acceder directamente desde la tienda: haga clic en el ícono de casa (🏠) en la esquina superior del panel para ir a la tienda, o en el logotipo para volver al panel.")
    sep(doc)
    body(doc, "Estructura del panel de administración:", bold=True)
    bullet(doc, "Barra lateral izquierda (Sidebar): menú de navegación con todas las secciones del panel. Puede colapsarla haciendo clic en el botón de las tres líneas (≡).")
    bullet(doc, "Barra superior (Topbar): muestra el nombre del administrador y acceso rápido a la tienda.")
    bullet(doc, "Área central: el contenido de la sección seleccionada.")
    bullet(doc, "Campana de notificaciones: alertas de stock bajo y pedidos nuevos.")
    sep(doc)

    h2(doc, "8.2 Dashboard (Tablero de mando)")
    body(doc,
        "El Dashboard es la primera pantalla que ve al entrar al panel. Muestra un "
        "resumen rápido del estado de la tienda en tiempo real."
    )
    body(doc, "Información que verá en el Dashboard:", bold=True)
    bullet(doc, "Total de ventas del día / mes / período.")
    bullet(doc, "Número de pedidos pendientes.")
    bullet(doc, "Número de productos con stock bajo.")
    bullet(doc, "Número de usuarios registrados.")
    bullet(doc, "Gráficas de ventas recientes.")
    bullet(doc, "Últimos pedidos realizados.")
    sep(doc)
    tip_box(doc,
        "Use el Dashboard para tener una visión rápida de la situación de la tienda "
        "cada vez que ingrese al panel.",
        "tip"
    )
    sep(doc)

    h2(doc, "8.3 Gestión de Pedidos")
    body(doc,
        "Desde esta sección puede ver y gestionar todos los pedidos realizados por los clientes."
    )
    body(doc, "¿Qué puede hacer en esta sección?", bold=True)
    bullet(doc, "Ver la lista completa de pedidos con fecha, cliente, total y estado.")
    bullet(doc, "Filtrar pedidos por estado: PENDIENTE, EN PREPARACIÓN, ENTREGADO, CANCELADO.")
    bullet(doc, "Buscar pedidos por número o por nombre del cliente.")
    bullet(doc, "Ver el detalle de cada pedido (productos, dirección, datos del cliente).")
    bullet(doc, "Actualizar el estado de un pedido.")
    bullet(doc, "Exportar la lista de pedidos a Excel.")
    sep(doc)
    body(doc, "¿Cómo cambiar el estado de un pedido?", bold=True)
    number_item(doc, 1, 'Haga clic en "Pedidos" en el menú lateral.')
    number_item(doc, 2, "Encuentre el pedido que desea actualizar.")
    number_item(doc, 3, "Haga clic en el pedido para ver su detalle, o busque el botón de actualizar estado directamente en la lista.")
    number_item(doc, 4, "Seleccione el nuevo estado del pedido: EN PREPARACIÓN, ENTREGADO o CANCELADO.")
    number_item(doc, 5, "Confirme el cambio.")
    sep(doc)
    tip_box(doc,
        "Actualice el estado de los pedidos con frecuencia para mantener informados "
        "a los clientes sobre el progreso de sus compras.",
        "tip"
    )
    sep(doc)

    h2(doc, "8.4 Gestión de Productos")
    body(doc,
        "Desde esta sección puede crear, editar y eliminar los productos de la tienda."
    )
    body(doc, "Indicadores (KPIs) que verá:", bold=True)
    bullet(doc, "Productos Totales: cantidad total de productos en la tienda.")
    bullet(doc, "Stock Bajo: productos con pocas unidades disponibles (menos de 5 unidades).")
    bullet(doc, "Sin Stock: productos agotados que no pueden ser vendidos.")
    sep(doc)

    body(doc, "Crear un nuevo producto:", bold=True)
    number_item(doc, 1, 'Haga clic en "Productos" en el menú lateral.')
    number_item(doc, 2, 'Haga clic en el botón "Nuevo Producto" (generalmente en la esquina superior derecha).')
    number_item(doc, 3, "Complete el formulario con los datos del producto:")
    bullet(doc, "Nombre del producto: nombre claro y descriptivo.", level=1)
    bullet(doc, "Descripción: información detallada del producto (características, usos, peso, etc.).", level=1)
    bullet(doc, "Precio: valor del producto en pesos colombianos (COP). Nota: el sistema trabaja internamente en centavos, así que para ingresar $25.000 deberá escribir 2500000.", level=1)
    bullet(doc, "Stock inicial: cuántas unidades hay disponibles.", level=1)
    bullet(doc, "Categoría: seleccione la categoría a la que pertenece el producto.", level=1)
    bullet(doc, "Imágenes: suba las fotos del producto (máximo 5 MB por imagen).", level=1)
    bullet(doc, "Estado (Activo/Inactivo): solo los productos activos son visibles para los clientes.", level=1)
    number_item(doc, 4, 'Haga clic en "Guardar producto".')
    sep(doc)

    body(doc, "Editar un producto existente:", bold=True)
    number_item(doc, 1, "Encuentre el producto en la lista.")
    number_item(doc, 2, 'Haga clic en el ícono de lápiz (✏) o en el botón "Editar".')
    number_item(doc, 3, "Modifique los campos que desea cambiar.")
    number_item(doc, 4, 'Haga clic en "Guardar cambios".')
    sep(doc)

    body(doc, "Desactivar / Eliminar un producto:", bold=True)
    number_item(doc, 1, "Encuentre el producto en la lista.")
    number_item(doc, 2, 'Haga clic en el ícono de basura (🗑) o en "Eliminar".')
    number_item(doc, 3, "El producto pasará a estado 'Inactivo' (desaparece de la tienda pero sus datos se conservan).")
    sep(doc)
    tip_box(doc,
        "Los productos eliminados no desaparecen permanentemente de la base de datos. "
        "Simplemente se marcan como inactivos. Esto preserva el historial de ventas.",
        "info"
    )
    tip_box(doc,
        "Asegúrese de subir imágenes de buena calidad (fondo blanco, bien iluminadas). "
        "Esto aumenta significativamente las ventas.",
        "tip"
    )
    sep(doc)

    h2(doc, "8.5 Control de Inventario")
    body(doc,
        "La sección de Inventario le permite ver y controlar el stock de todos "
        "los productos de forma centralizada."
    )
    body(doc, "¿Qué puede hacer?", bold=True)
    bullet(doc, "Ver el stock actual de todos los productos en una sola tabla.")
    bullet(doc, "Identificar rápidamente los productos agotados o con stock bajo.")
    bullet(doc, "Registrar movimientos de inventario: compras (ingreso de mercancía), ajustes, devoluciones.")
    bullet(doc, "Ver el historial completo de movimientos de un producto.")
    bullet(doc, "Exportar el inventario a Excel.")
    sep(doc)

    body(doc, "¿Cómo registrar un movimiento de inventario?", bold=True)
    body(doc, "Use esto cuando reciba mercancía nueva, haga un ajuste de inventario o registre una devolución:")
    number_item(doc, 1, 'Haga clic en "Inventario" en el menú lateral.')
    number_item(doc, 2, "Encuentre el producto cuyo stock desea modificar.")
    number_item(doc, 3, 'Haga clic en "Registrar movimiento" o en el ícono correspondiente.')
    number_item(doc, 4, "Seleccione el tipo de movimiento:")
    bullet(doc, "COMPRA: cuando recibe mercancía nueva (aumenta el stock).", level=1)
    bullet(doc, "AJUSTE: corrección manual del inventario.", level=1)
    bullet(doc, "DEVOLUCION: cuando un cliente devuelve un producto.", level=1)
    number_item(doc, 5, "Ingrese la cantidad y una descripción del movimiento.")
    number_item(doc, 6, "Confirme el registro.")
    sep(doc)
    tip_box(doc,
        "Las ventas descuentan el stock automáticamente cuando un cliente realiza un pedido. "
        "No es necesario hacer movimientos manuales por ventas.",
        "info"
    )
    sep(doc)

    h2(doc, "8.6 Categorías")
    body(doc,
        "Las categorías organizan los productos para que los clientes los encuentren "
        "más fácilmente. Puede crear, editar y organizar las categorías."
    )
    body(doc, "Crear una categoría:", bold=True)
    number_item(doc, 1, 'Haga clic en "Categorías" en el menú lateral.')
    number_item(doc, 2, 'Haga clic en "Nueva Categoría".')
    number_item(doc, 3, "Complete el formulario:")
    bullet(doc, "Nombre: nombre de la categoría (ej: 'Cafés Especiales').", level=1)
    bullet(doc, "Descripción: breve descripción de qué tipo de productos incluye.", level=1)
    bullet(doc, "Ícono: ícono visual que representa la categoría.", level=1)
    number_item(doc, 4, 'Haga clic en "Guardar".')
    sep(doc)
    tip_box(doc,
        "No puede eliminar una categoría que tenga productos asignados. "
        "Primero reasigne o elimine los productos de esa categoría.",
        "warning"
    )
    sep(doc)

    h2(doc, "8.7 Usuarios")
    body(doc,
        "Desde esta sección puede ver y gestionar las cuentas de todos los usuarios "
        "registrados en la plataforma."
    )
    body(doc, "¿Qué puede hacer?", bold=True)
    bullet(doc, "Ver la lista completa de usuarios (clientes y administradores).")
    bullet(doc, "Ver el correo electrónico, nombre, teléfono y rol de cada usuario.")
    bullet(doc, "Activar o desactivar cuentas de usuario.")
    bullet(doc, "Ver la fecha de registro de cada usuario.")
    sep(doc)
    tip_box(doc,
        "Para proteger la privacidad, las contraseñas de los usuarios nunca son "
        "visibles, ni siquiera para los administradores.",
        "info"
    )
    tip_box(doc,
        "Desactive una cuenta de usuario solo cuando sea estrictamente necesario "
        "(ej: actividad fraudulenta). El usuario no podrá iniciar sesión mientras esté inactivo.",
        "warning"
    )
    sep(doc)

    h2(doc, "8.8 Ofertas")
    body(doc,
        "La sección de Ofertas le permite crear y gestionar promociones y descuentos "
        "especiales en los productos."
    )
    body(doc, "Crear una oferta:", bold=True)
    number_item(doc, 1, 'Haga clic en "Ofertas" en el menú lateral.')
    number_item(doc, 2, 'Haga clic en "Nueva Oferta".')
    number_item(doc, 3, "Configure la oferta:")
    bullet(doc, "Producto: seleccione el producto al que quiere aplicar el descuento.", level=1)
    bullet(doc, "Porcentaje de descuento: ej: 15 (para 15% de descuento).", level=1)
    bullet(doc, "Fecha de inicio y fin: período en que estará activa la oferta.", level=1)
    bullet(doc, "Descripción: texto corto que verá el cliente (ej: '¡Oferta de temporada!').", level=1)
    number_item(doc, 4, 'Haga clic en "Guardar oferta".')
    sep(doc)
    tip_box(doc,
        "Las ofertas activas aparecen automáticamente en la sección pública 'Ofertas' "
        "de la tienda y el precio con descuento se muestra en la tarjeta del producto.",
        "tip"
    )
    sep(doc)

    h2(doc, "8.9 Reportes")
    body(doc,
        "La sección de Reportes le da información analítica sobre el desempeño de la tienda."
    )
    body(doc, "Tipos de reportes disponibles:", bold=True)
    tbl_rep = tbl_header(doc, ["Reporte", "¿Qué muestra?", "Utilidad"])
    for row in [
        ("Ventas",
         "Ingresos totales por período (día, semana, mes).",
         "Conocer cuánto dinero entró en la tienda."),
        ("Productos más vendidos",
         "Ranking de productos por cantidad de unidades vendidas.",
         "Identificar los productos estrella."),
        ("Baja Rotación",
         "Productos con pocas ventas en un período.",
         "Decidir qué productos promover o descontinuar."),
        ("Inventario",
         "Stock actual, productos agotados y stock bajo.",
         "Planear reposición de mercancía."),
    ]:
        add_row(tbl_rep, row)
    sep(doc)
    body(doc, "¿Cómo exportar un reporte?", bold=True)
    number_item(doc, 1, 'Haga clic en "Reportes" en el menú lateral.')
    number_item(doc, 2, "Seleccione el tipo de reporte que desea ver.")
    number_item(doc, 3, "Aplique los filtros de fecha si es necesario.")
    number_item(doc, 4, 'Haga clic en el botón "Exportar" (Excel o CSV).')
    number_item(doc, 5, "El archivo se descargará automáticamente en su computadora.")
    sep(doc)
    tip_box(doc,
        "Revise los reportes al menos una vez por semana para estar al tanto de "
        "la situación de ventas e inventario de la tienda.",
        "tip"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 9 — FAQ y SOLUCIÓN DE PROBLEMAS
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 9 — Preguntas Frecuentes y Solución de Problemas")

    preguntas = [
        ("¿Puedo comprar sin registrarme?",
         "Puede explorar la tienda sin registrarse, pero para agregar productos al carrito "
         "y realizar una compra necesita tener una cuenta e iniciar sesión. El registro es "
         "completamente gratuito."),

        ("¿Es seguro comprar en StoreHub?",
         "Sí. StoreHub utiliza tecnologías de seguridad modernas: HTTPS (conexión cifrada), "
         "JWT (tokens de autenticación seguros) y sus contraseñas se almacenan de forma cifrada. "
         "Nunca compartimos sus datos con terceros."),

        ("¿La página no carga o carga muy lento. ¿Qué hago?",
         "1. Verifique su conexión a internet.\n"
         "2. Intente recargar la página (F5 o Ctrl+R).\n"
         "3. Borre el caché del navegador: Ctrl+Shift+Del en Chrome/Firefox.\n"
         "4. Intente abrir la página en modo incógnito (Ctrl+Shift+N).\n"
         "5. Pruebe con otro navegador."),

        ("Olvidé mi contraseña y no puedo acceder. ¿Qué hago?",
         "Contacte al administrador de la tienda indicando su nombre y el correo con que se "
         "registró, para que pueda ayudarle a recuperar el acceso a su cuenta."),

        ("Agregué productos al carrito pero no aparecen. ¿Por qué?",
         "Esto puede ocurrir si: (a) No inició sesión antes de agregar al carrito, (b) Su "
         "sesión expiró. Inicie sesión nuevamente y los productos deberían seguir en su carrito. "
         "Si el problema persiste, intente recargar la página."),

        ("¿Por qué no puedo agregar más unidades de un producto?",
         "Cada producto tiene un stock limitado. Si el sistema no le permite agregar más "
         "unidades, significa que no hay suficiente inventario disponible. Puede intentar más "
         "adelante cuando el administrador reponga el stock."),

        ("¿Puedo modificar o cancelar un pedido ya realizado?",
         "Una vez confirmado, el pedido no puede ser modificado directamente por el cliente. "
         "Contacte al equipo de la tienda lo más pronto posible indicando su número de pedido "
         "y el cambio que necesita. Los pedidos en estado PENDIENTE pueden cancelarse más fácilmente."),

        ("¿Cómo sé cuándo llegará mi pedido?",
         "Puede hacer seguimiento desde la sección 'Mis Pedidos' (ver Sección 6). Cuando el "
         "estado cambie a 'ENTREGADO', significa que el pedido fue procesado. Los tiempos de "
         "entrega física dependen de los acuerdos con Distribuciones S y J."),

        ("La página muestra un error rojo en la parte superior. ¿Qué significa?",
         "Los mensajes de error generalmente explican el problema (ej: 'Correo o contraseña "
         "incorrectos', 'Producto agotado'). Léalo con atención y siga las instrucciones. "
         "Si el error es genérico ('Error del servidor'), intente recargar la página o "
         "contacte al administrador."),

        ("¿Puedo usar StoreHub desde mi celular?",
         "Sí. StoreHub está diseñado para funcionar correctamente en celulares y tabletas. "
         "Simplemente abra su navegador web (Chrome, Safari, Firefox) y acceda a la dirección "
         "de la tienda. No necesita descargar ninguna aplicación."),

        ("¿Cómo ingreso la clave de administrador? ¿Cuál es?",
         "Las credenciales de administrador son asignadas por el equipo técnico y no se "
         "comparten públicamente. Si necesita acceso de administrador, contacte a los "
         "desarrolladores del sistema."),

        ("¿Qué hago si veo 'Sesión expirada'?",
         "Su sesión expiró por inactividad (después de 1 hora). Simplemente inicie sesión "
         "nuevamente con su correo y contraseña. Sus datos y carrito se conservan."),
    ]

    for pregunta, respuesta in preguntas:
        h3(doc, pregunta)
        body(doc, respuesta)
        sep(doc)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECCIÓN 10 — GLOSARIO
    # ═══════════════════════════════════════════
    h1(doc, "SECCIÓN 10 — Glosario de Términos")
    body(doc,
        "Esta sección explica, en lenguaje sencillo, los términos técnicos que puede "
        "encontrar al usar StoreHub."
    )
    sep(doc)

    terminos = [
        ("API", "Sistema interno de comunicación entre el servidor y la aplicación. El usuario no interactúa directamente con ella."),
        ("Backend", "La parte del sistema que funciona 'detrás de escena': base de datos, lógica de negocio, seguridad. No es visible para el usuario."),
        ("Carrito de compras", "Espacio virtual donde se guardan temporalmente los productos que desea comprar antes de confirmar el pedido."),
        ("Caché", "Datos guardados temporalmente en su navegador para que las páginas carguen más rápido."),
        ("Checkout", "El proceso de finalizar una compra: seleccionar dirección, revisar el pedido y confirmarlo."),
        ("Cookie", "Pequeño archivo que el sitio web guarda en su navegador para recordar su sesión y preferencias."),
        ("Frontend", "La parte visual de la aplicación que usted ve e interactúa: botones, formularios, imágenes."),
        ("HTTPS", "Protocolo de conexión segura. El candado verde en la barra del navegador indica que la conexión es cifrada y segura."),
        ("JWT (Token)", "Código cifrado que el sistema usa para verificar que usted está correctamente autenticado. Se guarda en su navegador."),
        ("Soft delete", "Cuando un elemento es 'eliminado' pero en realidad solo se oculta. Los datos se conservan en la base de datos para mantener el historial."),
        ("Stock", "Cantidad de unidades disponibles de un producto en el inventario."),
        ("Stock bajo", "Cuando quedan pocas unidades de un producto (generalmente menos de 5 unidades)."),
        ("Usuario ADMIN", "Perfil con permisos completos para gestionar la tienda: productos, pedidos, usuarios, etc."),
        ("Usuario CLIENTE", "Perfil estándar que permite comprar en la tienda."),
        ("Stock bajo", "Alerta que aparece cuando un producto tiene pocas unidades disponibles."),
        ("Soft delete / Inactivo", "Cuando un producto u usuario se 'elimina' de la vista pero sus datos se conservan en el sistema."),
    ]

    tbl_glos = tbl_header(doc, ["Término", "Definición"])
    seen = set()
    for termino, definicion in terminos:
        if termino not in seen:
            add_row(tbl_glos, [termino, definicion])
            seen.add(termino)
    sep(doc)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # CONTRAPORTADA / INFO FINAL
    # ═══════════════════════════════════════════
    h1(doc, "Información del Documento")
    tbl_info = tbl_header(doc, ["Campo", "Valor"])
    for row in [
        ("Documento",       "Manual de Usuario — StoreHub"),
        ("Versión",         "1.0.0"),
        ("Fecha",           "04 de Marzo de 2026"),
        ("Plataforma",      "https://inventory-2-sewi.onrender.com"),
        ("API Backend",     "https://inventory-1-jkh2.onrender.com/api/v1"),
        ("Cliente",         "Julián Alberto Garzón García / Distribuciones S y J"),
        ("Institución",     "SENA — Servicio Nacional de Aprendizaje"),
        ("Autores",         "Sebastián Gallego V., Santiago Sánchez S., Duván Ochoa H., Ricky Lotero S."),
        ("Generado",        datetime.datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]:
        add_row(tbl_info, row)

    return doc


# ─── main ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    output = r"c:\Users\Equipo\Desktop\inventory app\MANUAL_USUARIO_StoreHub.docx"
    doc = build()
    doc.save(output)
    print(f"[OK] Manual de usuario generado: {output}")
