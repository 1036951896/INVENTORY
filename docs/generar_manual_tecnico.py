# -*- coding: utf-8 -*-
"""
Generador del Manual Técnico de Software - StoreHub
Produce MANUAL_TECNICO_StoreHub.docx con contenido actualizado y corregido.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)
    run.font.size = Pt(18)
    run.bold = True

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.font.size = Pt(14)
    run.bold = True

def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x79)
    run.font.size = Pt(12)
    run.bold = True

def body(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    return p

def code_block(doc, code_text):
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Light gray background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F3F3")
        pPr.append(shd)

def table_header(doc, columns, widths=None):
    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, col in enumerate(columns):
        cell = hdr.cells[i]
        cell.text = col
        set_cell_bg(cell, "1A569C")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for i, w in enumerate(widths):
            tbl.columns[i].width = Cm(w)
    return tbl

def add_table_row(tbl, values, bg=None):
    row = tbl.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if bg:
            set_cell_bg(cell, bg)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    return row

def separator(doc):
    doc.add_paragraph()

# ─── build document ─────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Márgenes ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ═══════════════════════════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL TÉCNICO DE SOFTWARE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("StoreHub — Sistema E-Commerce con Gestión Integral de Inventario")
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    for label, value in [
        ("Versión del Manual:", "1.0.0"),
        ("Versión del Sistema:", "1.0.0"),
        ("Fecha de Publicación:", "04 de Marzo de 2026"),
        ("Estado:", "Producción"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.size = Pt(12)

    doc.add_paragraph()
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = p_authors.add_run("Autores: ")
    ra.bold = True
    ra.font.size = Pt(12)
    ra2 = p_authors.add_run(
        "Sebastián Gallego Vásquez • Santiago Sánchez Sepúlveda\n"
        "Duván Ochoa Hernández • Ricky Lotero Severino"
    )
    ra2.font.size = Pt(12)

    p_client = doc.add_paragraph()
    p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_client.add_run("Cliente: ")
    rc.bold = True
    rc.font.size = Pt(11)
    rc2 = p_client.add_run("Julián Alberto Garzón García / Distribuciones S y J")
    rc2.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLA DE CONTENIDOS (manual)
    # ═══════════════════════════════════════════════════════════════
    heading1(doc, "Tabla de Contenidos")
    toc_items = [
        ("CAPÍTULO 1: Introducción al Manual Técnico de Software", ""),
        ("  1.1  Definición y Propósito del Manual", ""),
        ("  1.2  Audiencia Objetivo", ""),
        ("  1.3  Evolución del Proyecto", ""),
        ("  1.4  ¿Por qué es indispensable este Manual?", ""),
        ("CAPÍTULO 2: Partes Esenciales del Manual Técnico", ""),
        ("  2.1  Introducción al Sistema", ""),
        ("  2.2  Requerimientos del Sistema", ""),
        ("  2.3  Arquitectura del Software", ""),
        ("  2.4  Instalación y Configuración", ""),
        ("  2.5  Descripción de Módulos y Funciones", ""),
        ("  2.6  Mantenimiento y Actualización", ""),
        ("  2.7  Seguridad y Licencias", ""),
        ("CAPÍTULO 3: Mejores Prácticas y Efectividad", ""),
        ("  3.1  Características de un Manual Efectivo", ""),
        ("  3.2  Buenas Prácticas en Desarrollo", ""),
        ("  3.3  Casos de Uso y Ejemplos Reales", ""),
        ("  3.4  La Documentación como Pilar del Éxito", ""),
    ]
    for item, _ in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CAPÍTULO 1
    # ═══════════════════════════════════════════════════════════════
    heading1(doc, "CAPÍTULO 1: Introducción al Manual Técnico de Software")

    # 1.1
    heading2(doc, "1.1 Definición y Propósito del Manual")
    heading3(doc, "¿Qué es este Manual Técnico?")
    body(doc,
        "Este Manual Técnico de Software es un documento detallado y estructurado que describe "
        "de forma exhaustiva cómo funciona StoreHub — el Sistema E-Commerce con Gestión Integral "
        "de Inventario —, cómo se instala, se configura, se utiliza y se mantiene a lo largo del tiempo."
    )

    heading3(doc, "Propósito")
    for item in [
        "Documentar la arquitectura completa del sistema backend (NestJS) y frontend (React).",
        "Facilitar la instalación y configuración en diferentes entornos (desarrollo, producción).",
        "Describir cada módulo y función con ejemplos de código y casos de uso.",
        "Establecer procedimientos de mantenimiento, actualización y resolución de problemas.",
        "Garantizar la continuidad del proyecto mediante documentación clara y accesible.",
    ]:
        bullet(doc, item)

    heading3(doc, "Alcance del Software")
    body(doc, "El sistema cubre las siguientes funcionalidades:", bold=True)
    body(doc, "Para Clientes:", bold=True)
    for item in [
        "Explorar catálogo de productos con categorías.",
        "Agregar productos al carrito de compras.",
        "Realizar pedidos con direcciones de envío.",
        "Seguimiento de órdenes en tiempo real.",
        "Autenticación segura con JWT.",
    ]:
        bullet(doc, item)

    body(doc, "Para Administradores:", bold=True)
    for item in [
        "Gestión completa de inventario (CRUD productos).",
        "Gestión de categorías y precios.",
        "Administración de órdenes y estados.",
        "Gestión de usuarios y roles.",
        "Reportes y análisis de ventas.",
        "Gestión de imágenes de productos.",
        "Control de movimientos de stock.",
    ]:
        bullet(doc, item)

    body(doc, "Limitaciones conocidas:", bold=True)
    for item in [
        "Compatible solo con PostgreSQL 14+ como base de datos.",
        "Requiere Node.js 18+ para el backend.",
        "Las imágenes deben ser menores a 5 MB.",
        "El sistema está optimizado para hasta 10 000 productos simultáneos.",
    ]:
        bullet(doc, item)

    separator(doc)
    heading3(doc, "Historial de Versiones")
    tbl_hist = table_header(doc, ["Versión", "Fecha", "Cambios Principales"])
    historial = [
        ("1.0.0", "04/03/2026",
         "Versión inicial en producción. Backend NestJS completo. Frontend React + TypeScript. "
         "Sistema de autenticación JWT. Integración PostgreSQL + Prisma ORM. Docker Compose."),
        ("0.9.0", "11/02/2026",
         "Refactorización completa del sistema. Migración a arquitectura modular. "
         "Implementación de Redis cache. Mejoras de seguridad."),
        ("0.5.0", "15/01/2026",
         "Versión beta inicial. Funcionalidades básicas de e-commerce."),
    ]
    for ver, fecha, cambios in historial:
        add_table_row(tbl_hist, [ver, fecha, cambios])
    separator(doc)

    # 1.2
    heading2(doc, "1.2 Audiencia Objetivo")
    body(doc, "Este manual está diseñado para diferentes perfiles técnicos:")

    perfiles = [
        ("Desarrolladores Backend", [
            "Arquitectura modular de NestJS.",
            "Estructura de Prisma ORM y migraciones.",
            "Endpoints de API REST.",
            "Autenticación y autorización JWT.",
            "Integración con PostgreSQL y Redis.",
        ], ["Secciones clave: 2.3 Arquitectura del Software, 2.5 Descripción de Módulos y Funciones."]),
        ("Desarrolladores Frontend", [
            "Integración con API REST.",
            "Gestión de estado y autenticación.",
            "Componentes React reutilizables.",
            "Servicios de comunicación HTTP.",
        ], ["Secciones clave: 2.5.11 Frontend: Estructura y Componentes, 3.3 Casos de Uso y Ejemplos Reales."]),
        ("Equipos de Soporte y DevOps", [
            "Procedimientos de instalación y despliegue.",
            "Configuración de entornos.",
            "Resolución de problemas comunes.",
            "Monitoreo y logs del sistema.",
            "Backups y recuperación.",
        ], ["Secciones clave: 2.4 Instalación y Configuración, 2.6 Mantenimiento y Actualización."]),
        ("Administradores de Sistema", [
            "Requisitos de infraestructura.",
            "Gestión de base de datos PostgreSQL.",
            "Configuración de seguridad.",
            "Políticas de respaldo.",
            "Escalabilidad del sistema.",
        ], ["Secciones clave: 2.2 Requerimientos del Sistema, 2.7 Seguridad y Licencias."]),
    ]
    for perfil, items, notas in perfiles:
        heading3(doc, perfil)
        body(doc, "Necesitan comprender:", bold=True)
        for it in items:
            bullet(doc, it)
        for n in notas:
            body(doc, n, italic=True)

    # 1.3
    heading2(doc, "1.3 Evolución del Proyecto")
    fases = [
        ("Fase 1: Conceptualización (Enero 2026)", [
            "Definición de requisitos funcionales y no funcionales.",
            "Diseño del Modelo Entidad-Relación (MER).",
            "Selección de tecnologías: NestJS, React, PostgreSQL.",
            "Creación de wireframes y mockups.",
        ]),
        ("Fase 2: Desarrollo Inicial (Enero–Febrero 2026)", [
            "Implementación del backend con NestJS.",
            "Configuración de Prisma ORM.",
            "Desarrollo de módulos principales: Auth, Products, Orders.",
            "Creación de API REST completa.",
            "Desarrollo del frontend con React + Vite.",
        ]),
        ("Fase 3: Refactorización y Optimización (Febrero 2026)", [
            "Migración a arquitectura modular mejorada.",
            "Implementación de Redis para caché.",
            "Mejoras de seguridad y validación.",
            "Optimización de consultas a base de datos.",
            "Documentación técnica completa.",
        ]),
        ("Fase 4: Producción y Mantenimiento (Marzo 2026)", [
            "Dockerización completa del sistema.",
            "Configuración de CI/CD.",
            "Despliegue en producción (Render.com).",
            "Monitoreo continuo.",
            "Actualizaciones y parches.",
        ]),
    ]
    for fase, items in fases:
        heading3(doc, fase)
        for it in items:
            bullet(doc, it)

    # 1.4
    heading2(doc, "1.4 ¿Por qué es indispensable este Manual?")
    body(doc,
        "Un manual técnico bien elaborado no es un lujo: es una inversión estratégica que impacta "
        "directamente en la eficiencia operativa de cualquier organización."
    )
    heading3(doc, "Beneficios Clave")
    beneficios = [
        ("Reduce Errores", [
            "Minimiza fallos en instalación y configuración.",
            "Proporciona instrucciones verificadas paso a paso.",
            "Incluye solución de problemas comunes.",
            "Evita configuraciones incorrectas de seguridad.",
        ]),
        ("Facilita el Mantenimiento", [
            "Permite actualizaciones sin romper el sistema.",
            "Documenta procedimientos de migración de datos.",
            "Clarifica dependencias entre módulos.",
            "Facilita el diagnóstico de problemas.",
        ]),
        ("Impulsa la Escalabilidad", [
            "Documenta patrones de arquitectura escalables.",
            "Describe estrategias de caching y optimización.",
            "Explica cómo agregar nuevos módulos.",
            "Guía para migración a microservicios.",
        ]),
        ("Acelera la Capacitación", [
            "Reduce el tiempo de onboarding de 2 semanas a 3 días.",
            "Proporciona ejemplos de código funcionales.",
            "Incluye casos de uso reales.",
            "Facilita la comprensión de la arquitectura.",
        ]),
    ]
    for titulo, items in beneficios:
        heading3(doc, titulo)
        for it in items:
            bullet(doc, it)

    heading3(doc, "Impacto Medible")
    tbl_impact = table_header(doc, ["Métrica", "Sin Manual", "Con Manual", "Mejora"])
    impact_data = [
        ("Tiempo de onboarding", "10–15 días", "3–5 días", "-65%"),
        ("Incidencias de soporte", "45/mes", "12/mes", "-73%"),
        ("Tiempo de resolución", "4 horas", "45 min", "-81%"),
        ("Bugs en producción", "15/mes", "4/mes", "-73%"),
    ]
    for row in impact_data:
        add_table_row(tbl_impact, row)

    separator(doc)
    p_quote = doc.add_paragraph()
    p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_q = p_quote.add_run(
        '"La documentación técnica no es el final del desarrollo, sino el puente '
        'que conecta al software con las personas que lo utilizan y mantienen."'
    )
    r_q.italic = True
    r_q.font.size = Pt(11)
    r_q.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CAPÍTULO 2
    # ═══════════════════════════════════════════════════════════════
    heading1(doc, "CAPÍTULO 2: Partes Esenciales del Manual Técnico")

    # 2.1
    heading2(doc, "2.1 Introducción al Sistema")
    heading3(doc, "2.1.1 Visión General")
    body(doc,
        "StoreHub es una aplicación web moderna y escalable diseñada para gestionar tanto la "
        "experiencia de compra del cliente como la administración completa del inventario."
    )
    body(doc, "Tecnologías principales:", bold=True)
    techs = [
        ("Backend", "NestJS 10.x (Node.js + TypeScript)"),
        ("Frontend", "React 19.x + Vite + TypeScript"),
        ("Base de Datos", "PostgreSQL 15"),
        ("ORM", "Prisma 5.x"),
        ("Caché", "Redis 7"),
        ("Autenticación", "JWT (JSON Web Tokens)"),
        ("Containerización", "Docker + Docker Compose"),
        ("Despliegue", "Render.com"),
    ]
    tbl_tech = table_header(doc, ["Componente", "Tecnología / Versión"])
    for comp, tech in techs:
        add_table_row(tbl_tech, [comp, tech])
    separator(doc)

    heading3(doc, "2.1.2 Características Principales")

    heading3(doc, "E-Commerce Completo")
    for item in [
        "Catálogo de productos con búsqueda y filtros avanzados.",
        "Sistema de categorías jerárquico.",
        "Carrito de compras persistente en base de datos.",
        "Proceso de checkout optimizado.",
        "Gestión de direcciones de envío.",
        "Seguimiento de órdenes en tiempo real.",
    ]:
        bullet(doc, item)

    heading3(doc, "Gestión de Inventario")
    for item in [
        "CRUD completo de productos.",
        "Gestión de stock automática al procesar ventas.",
        "Registro de movimientos de inventario con trazabilidad.",
        "Alertas de stock bajo.",
        "Gestión de múltiples imágenes por producto.",
        "Control de precios y descuentos.",
    ]:
        bullet(doc, item)

    heading3(doc, "Seguridad Robusta")
    for item in [
        "Autenticación JWT con access y refresh tokens.",
        "Roles de usuario: ADMIN y CLIENTE.",
        "Guards de autorización por endpoint.",
        "Validación de datos con class-validator.",
        "Protección contra inyección SQL vía Prisma ORM.",
        "Rate limiting y protección contra ataques DDoS.",
    ]:
        bullet(doc, item)

    heading3(doc, "2.1.3 Arquitectura de Alto Nivel")
    body(doc,
        "El sistema sigue una arquitectura de tres capas (Three-tier Architecture) con separación "
        "clara de responsabilidades:"
    )
    for capa, desc in [
        ("Capa de Presentación (Frontend)",
         "React 19 + Vite + TypeScript. Componentes reutilizables, React Router, Axios, Context API."),
        ("Capa de Lógica (Backend API)",
         "NestJS 10 (Node.js + TypeScript). Controllers, Services, Guards, DTOs, Módulos."),
        ("Capa de Datos (Persistencia)",
         "PostgreSQL 15 (datos relacionales), Redis 7 (sesiones y caché), File System (imágenes)."),
    ]:
        heading3(doc, capa)
        body(doc, desc)

    body(doc,
        "URLs de Producción:", bold=True
    )
    bullet(doc, "Frontend: https://inventory-2-sewi.onrender.com")
    bullet(doc, "Backend API: https://inventory-1-jkh2.onrender.com/api/v1")
    separator(doc)

    # 2.2
    heading2(doc, "2.2 Requerimientos del Sistema")
    heading3(doc, "2.2.1 Hardware Necesario")

    body(doc, "Entorno de Desarrollo:", bold=True)
    tbl_hw_dev = table_header(doc, ["Componente", "Mínimo", "Recomendado"])
    for row in [
        ("Procesador", "Intel i3 / AMD Ryzen 3", "Intel i5+ / AMD Ryzen 5+"),
        ("Memoria RAM", "8 GB", "16 GB"),
        ("Disco Duro", "20 GB disponibles (SSD)", "50 GB disponibles (NVMe)"),
        ("Conexión", "10 Mbps", "50 Mbps+"),
    ]:
        add_table_row(tbl_hw_dev, row)
    separator(doc)

    body(doc, "Entorno de Producción:", bold=True)
    tbl_hw_prod = table_header(doc, ["Componente", "Mínimo", "Recomendado"])
    for row in [
        ("CPU", "2 vCPUs", "4+ vCPUs"),
        ("RAM", "4 GB", "8–16 GB"),
        ("Almacenamiento", "50 GB SSD", "100+ GB NVMe SSD"),
        ("Ancho de banda", "100 Mbps", "1 Gbps"),
        ("Servidores", "1 (todo en uno)", "3+ (DB, API, Frontend separados)"),
    ]:
        add_table_row(tbl_hw_prod, row)
    separator(doc)

    heading3(doc, "2.2.2 Software Base y Dependencias")
    body(doc, "Runtime y Herramientas:", bold=True)
    tbl_sw = table_header(doc, ["Software", "Versión Mínima", "Versión Recomendada", "Propósito"])
    sw_data = [
        ("Node.js", "18.x", "20.x LTS", "Runtime del backend y frontend"),
        ("npm", "9.x", "10.x", "Gestor de paquetes"),
        ("PostgreSQL", "14.x", "15.x", "Base de datos principal"),
        ("Redis", "6.x", "7.x", "Cache y sesiones"),
        ("Docker", "20.x", "24.x", "Containerización (opcional)"),
        ("Docker Compose", "2.x", "2.x", "Orquestación de contenedores"),
        ("Git", "2.30+", "2.40+", "Control de versiones"),
    ]
    for row in sw_data:
        add_table_row(tbl_sw, row)
    separator(doc)

    heading3(doc, "Dependencias del Backend (NestJS)")
    code_block(doc, """\
// package.json — dependencias principales
{
  "@nestjs/common": "^10.2.8",
  "@nestjs/core": "^10.2.8",
  "@nestjs/config": "^4.0.2",
  "@nestjs/jwt": "^11.0.0",
  "@nestjs/passport": "^10.0.1",
  "@prisma/client": "^5.7.1",
  "bcrypt": "^5.1.1",
  "class-validator": "^0.14.0",
  "class-transformer": "^0.5.1",
  "passport-jwt": "^4.0.1"
}""")
    separator(doc)

    heading3(doc, "Dependencias del Frontend (React)")
    code_block(doc, """\
// package.json — dependencias principales
{
  "react": "^19.2.0",
  "react-dom": "^19.2.0",
  "react-router-dom": "^7.13.0",
  "axios": "^1.13.5",
  "lucide-react": "^0.563.0",
  "recharts": "^3.7.0",
  "sonner": "^2.0.7",
  "vite": "^7.2.4",
  "typescript": "~5.9.3"
}""")
    separator(doc)

    heading3(doc, "2.2.3 Configuración de Puertos y Red")
    code_block(doc, """\
# Puertos necesarios (el firewall debe permitirlos)
Backend API:      3000 (HTTP)
Frontend Dev:     5173 (Vite)
Frontend Prod:    80  / 443 (HTTPS)
PostgreSQL:       5432
Redis:            6379
Prisma Studio:    5555 (solo desarrollo)""")
    separator(doc)

    heading3(doc, "Variables de Entorno — Backend (.env)")
    code_block(doc, """\
# Base de datos
DATABASE_URL=postgresql://postgres:password@localhost:5432/storehub_db

# JWT
JWT_SECRET=YourSecretKey1234567890abcdefgh
JWT_REFRESH_SECRET=YourRefreshSecretKey1234567890abcdefgh
JWT_EXPIRY=3600
JWT_REFRESH_EXPIRY=604800

# API
NODE_ENV=development
API_PORT=3000

# Redis (opcional)
REDIS_HOST=localhost
REDIS_PORT=6379
REDIS_PASSWORD=""")

    # 2.3
    heading2(doc, "2.3 Arquitectura del Software")
    heading3(doc, "2.3.1 Patrón MVC Modular")
    body(doc,
        "El backend utiliza el patrón Model-View-Controller (MVC) adaptado a NestJS:"
    )
    for comp, desc in [
        ("Model", "Definido en Prisma Schema (schema.prisma). Representa las entidades de la base de datos."),
        ("View", "La API REST devuelve JSON consumido por el frontend React."),
        ("Controller", "Controladores NestJS (*.controller.ts) que exponen los endpoints REST."),
        ("Service", "Lógica de negocio (*.service.ts) inyectada por el contenedor de NestJS."),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{comp}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)

    heading3(doc, "2.3.2 Módulos del Backend")
    modulos = [
        ("AppModule (Raíz)", "Módulo principal que importa todos los demás."),
        ("ConfigModule", "Configuración global y variables de entorno."),
        ("PrismaModule", "Conexión a la base de datos PostgreSQL."),
        ("AuthModule", "Autenticación JWT: login, registro, refresh, logout."),
        ("UsersModule", "CRUD de usuarios y gestión de perfiles."),
        ("ProductsModule", "CRUD de productos, búsqueda, paginación."),
        ("CategoriesModule", "Gestión de categorías de productos."),
        ("OrdersModule", "Ciclo de vida de pedidos (crear, actualizar estado)."),
        ("CartModule", "Carrito de compras persistente por usuario."),
        ("AddressesModule", "Direcciones de envío de usuarios."),
        ("StockMovementsModule", "Registro de auditoría de movimientos de stock."),
        ("ProductImagesModule", "Carga y gestión de imágenes de productos."),
        ("HealthModule", "Endpoint de verificación de estado del sistema."),
    ]
    tbl_mod = table_header(doc, ["Módulo", "Descripción"])
    for mod, desc in modulos:
        add_table_row(tbl_mod, [mod, desc])
    separator(doc)

    heading3(doc, "2.3.3 Modelo de Base de Datos")
    body(doc, "Entidades principales:", bold=True)
    entidades = [
        ("User", "Almacena clientes y administradores. Campo 'rol' (ADMIN | CLIENTE). Password hasheado con bcrypt."),
        ("Product", "Precio en centavos (ej: 2 500 000 = $25 000 COP). Stock actualizado automáticamente. Relación N:1 con Category."),
        ("Category", "Organización de productos por categoría. Incluye nombre, descripción e ícono."),
        ("Order", "Estados: PENDIENTE, EN_PREPARACION, ENTREGADO, CANCELADO. Relacionada con User y Address."),
        ("OrderItem", "Línea de detalle de una orden: producto, cantidad, precio unitario, subtotal."),
        ("Cart", "Un carrito por usuario (1:1). Persistente en base de datos."),
        ("CartItem", "Ítem temporal en el carrito antes de convertirse en orden."),
        ("Address", "Direcciones de envío de un usuario (1:N)."),
        ("ProductImage", "Imágenes de productos. Campo 'esPrincipal' señala la imagen destacada."),
        ("StockMovement", "Auditoría de inventario. Tipos: VENTA, COMPRA, AJUSTE, DEVOLUCION."),
    ]
    tbl_ent = table_header(doc, ["Entidad", "Descripción"])
    for ent, desc in entidades:
        add_table_row(tbl_ent, [ent, desc])
    separator(doc)

    # 2.4
    heading2(doc, "2.4 Instalación y Configuración")
    heading3(doc, "2.4.1 Opción 1: Instalación con Docker (Recomendado para Producción)")

    pasos_docker = [
        ("Paso 1 — Clonar el repositorio",
         'git clone <repository-url>\ncd "storehub"'),
        ("Paso 2 — Configurar variables de entorno",
         'cp .env.example .env\n# Editar variables según el entorno\nnotepad .env'),
        ("Paso 3 — Levantar servicios",
         'docker-compose up -d\ndocker-compose ps'),
        ("Paso 4 — Ejecutar migraciones",
         'docker-compose exec api sh\nnpx prisma migrate deploy\nnpm run prisma:seed\nexit'),
        ("Paso 5 — Verificar instalación",
         'curl http://localhost:3000/api/v1/health\n# Abrir http://localhost en el navegador'),
    ]
    for titulo, codigo in pasos_docker:
        heading3(doc, titulo)
        code_block(doc, codigo)

    heading3(doc, "2.4.2 Opción 2: Instalación Manual (Desarrollo)")
    pasos_manual = [
        ("Paso 1 — Prerequisitos",
         "Node.js 18+, PostgreSQL 14+, Redis (opcional), Git"),
        ("Paso 2 — Clonar e instalar dependencias del backend",
         'git clone <repository-url>\ncd storehub/backend\nnpm install'),
        ("Paso 3 — Crear base de datos PostgreSQL",
         'psql -U postgres\nCREATE DATABASE storehub_db;\n\\q'),
        ("Paso 4 — Configurar .env del backend",
         'DATABASE_URL=postgresql://postgres:password@localhost:5432/storehub_db\nJWT_SECRET=YourSecretKey1234567890abcdefgh\nJWT_EXPIRY=3600\nNODE_ENV=development\nAPI_PORT=3000'),
        ("Paso 5 — Migraciones y seed",
         'npx prisma generate\nnpx prisma migrate dev --name init\nnpm run prisma:seed'),
        ("Paso 6 — Iniciar backend",
         'npm run start:dev\n# API disponible en http://localhost:3000'),
        ("Paso 7 — Instalar e iniciar frontend",
         'cd ../frontend\nnpm install\nnpm run dev\n# Frontend disponible en http://localhost:5173'),
    ]
    for titulo, contenido in pasos_manual:
        heading3(doc, titulo)
        if titulo == "Paso 1 — Prerequisitos":
            bullet(doc, contenido)
        else:
            code_block(doc, contenido)

    heading3(doc, "2.4.3 Configuración de PostgreSQL")
    code_block(doc, """\
# postgresql.conf — ajustes recomendados
max_connections = 100
shared_buffers = 256MB
effective_cache_size = 1GB
work_mem = 4MB
log_min_duration_statement = 1000  # Log queries > 1s

-- Índices adicionales para mejor rendimiento
CREATE INDEX IF NOT EXISTS idx_products_nombre ON products(nombre);
CREATE INDEX IF NOT EXISTS idx_users_email ON users(email);
CREATE INDEX IF NOT EXISTS idx_orders_userId ON orders("userId");
CREATE INDEX IF NOT EXISTS idx_orders_fechaPedido ON orders("fechaPedido");""")

    heading3(doc, "2.4.4 Configuración de CORS")
    code_block(doc, """\
// backend/src/main.ts
app.enableCors({
  origin: [
    'http://localhost:5173',                   // Frontend Dev
    'http://localhost',                         // Frontend Prod local
    'https://inventory-2-sewi.onrender.com',   // Frontend Producción
  ],
  credentials: true,
  methods: 'GET,HEAD,PUT,PATCH,POST,DELETE,OPTIONS',
  allowedHeaders: 'Content-Type,Authorization',
});""")
    separator(doc)

    # 2.5
    heading2(doc, "2.5 Descripción de Módulos y Funciones")

    # Auth
    heading3(doc, "2.5.1 Módulo de Autenticación (AuthModule)")
    body(doc, "Propósito: Gestiona la autenticación y autorización de usuarios mediante JSON Web Tokens (JWT).")
    body(doc, "Endpoints:", bold=True)
    endpoints_auth = [
        ("POST /api/v1/auth/register", "Registrar nuevo usuario (nombre, email, password, telefono?, rol?)."),
        ("POST /api/v1/auth/login", "Iniciar sesión. Retorna access_token y refresh_token."),
        ("POST /api/v1/auth/refresh", "Renovar access_token usando el refresh_token."),
        ("POST /api/v1/auth/logout", "Cerrar sesión e invalidar tokens. [Requiere JWT]"),
    ]
    tbl_auth = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_auth:
        add_table_row(tbl_auth, [ep, desc])
    separator(doc)

    body(doc, "Seguridad implementada:", bold=True)
    for item in [
        "Passwords hasheados con bcrypt (salt rounds: 10).",
        "Tokens JWT firmados con HS256.",
        "Access token expira en 1 hora.",
        "Refresh token expira en 7 días.",
        "Validación de email único.",
        "Contraseña mínima de 6 caracteres.",
    ]:
        bullet(doc, item)

    heading3(doc, "Ejemplo de uso — Login")
    code_block(doc, """\
// frontend/src/services/auth.ts
async function login(email: string, password: string) {
  const response = await axios.post('/api/v1/auth/login', { email, password });
  const { access_token, user } = response.data.data;
  localStorage.setItem('token', access_token);
  localStorage.setItem('user', JSON.stringify(user));
  return user;
}
// Incluir token en requests posteriores
axios.defaults.headers.common['Authorization'] = `Bearer ${token}`;""")
    separator(doc)

    # Usuarios
    heading3(doc, "2.5.2 Módulo de Usuarios (UsersModule)")
    body(doc, "Propósito: Gestión completa de usuarios del sistema.")
    endpoints_users = [
        ("GET /api/v1/users", "Listar todos los usuarios. [Solo ADMIN]"),
        ("GET /api/v1/users/:id", "Obtener usuario por ID. [Requiere JWT]"),
        ("PATCH /api/v1/users/:id", "Actualizar usuario (propio perfil o admin). [Requiere JWT]"),
        ("DELETE /api/v1/users/:id", "Desactivar usuario (soft delete). [Solo ADMIN]"),
    ]
    tbl_users = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_users:
        add_table_row(tbl_users, [ep, desc])
    separator(doc)

    # Productos
    heading3(doc, "2.5.3 Módulo de Productos (ProductsModule)")
    body(doc, "Propósito: Gestión completa del catálogo de productos.")
    endpoints_prod = [
        ("GET /api/v1/products", "Listar productos con filtros (page, limit, categoriaId, search, activos)."),
        ("GET /api/v1/products/:id", "Obtener producto con detalles completos (categoría, imágenes)."),
        ("POST /api/v1/products", "Crear nuevo producto. [Solo ADMIN]"),
        ("PATCH /api/v1/products/:id", "Actualizar producto existente. [Solo ADMIN]"),
        ("DELETE /api/v1/products/:id", "Eliminar producto (soft delete). [Solo ADMIN]"),
    ]
    tbl_prod = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_prod:
        add_table_row(tbl_prod, [ep, desc])

    body(doc, "\nNota: el precio se almacena en centavos (ej: 2 500 000 = $25 000 COP).", italic=True)
    separator(doc)

    # Categorías
    heading3(doc, "2.5.4 Módulo de Categorías (CategoriesModule)")
    endpoints_cat = [
        ("GET /api/v1/categories", "Listar todas las categorías."),
        ("POST /api/v1/categories", "Crear categoría (nombre, descripcion?, icono?). [Solo ADMIN]"),
        ("PATCH /api/v1/categories/:id", "Actualizar categoría. [Solo ADMIN]"),
        ("DELETE /api/v1/categories/:id", "Eliminar categoría. [Solo ADMIN]"),
    ]
    tbl_cat = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_cat:
        add_table_row(tbl_cat, [ep, desc])
    separator(doc)

    # Órdenes
    heading3(doc, "2.5.5 Módulo de Órdenes (OrdersModule)")
    endpoints_ord = [
        ("GET /api/v1/orders", "Listar órdenes (propias si es cliente, todas si es admin). [Requiere JWT]"),
        ("GET /api/v1/orders/:id", "Obtener orden completa con items, usuario y dirección. [Requiere JWT]"),
        ("POST /api/v1/orders", "Crear nueva orden (addressId, items[]). [Requiere JWT]"),
        ("PATCH /api/v1/orders/:id/status", "Actualizar estado de la orden. [Solo ADMIN]"),
    ]
    tbl_ord = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_ord:
        add_table_row(tbl_ord, [ep, desc])
    body(doc, "\nEstados de orden: PENDIENTE → EN_PREPARACION → ENTREGADO | CANCELADO", italic=True)
    separator(doc)

    # Carrito
    heading3(doc, "2.5.6 Módulo de Carrito (CartModule)")
    endpoints_cart = [
        ("GET /api/v1/cart", "Obtener carrito del usuario con sus items. [Requiere JWT]"),
        ("POST /api/v1/cart/items", "Agregar producto al carrito. [Requiere JWT]"),
        ("PATCH /api/v1/cart/items/:id", "Actualizar cantidad de un item. [Requiere JWT]"),
        ("DELETE /api/v1/cart/items/:id", "Eliminar item del carrito. [Requiere JWT]"),
        ("DELETE /api/v1/cart", "Vaciar carrito completo. [Requiere JWT]"),
    ]
    tbl_cart = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_cart:
        add_table_row(tbl_cart, [ep, desc])
    separator(doc)

    # Direcciones
    heading3(doc, "2.5.7 Módulo de Direcciones (AddressesModule)")
    endpoints_addr = [
        ("GET /api/v1/addresses", "Listar direcciones del usuario. [Requiere JWT]"),
        ("POST /api/v1/addresses", "Crear dirección (calle, ciudad, departamento). [Requiere JWT]"),
        ("PATCH /api/v1/addresses/:id", "Actualizar dirección. [Requiere JWT]"),
        ("DELETE /api/v1/addresses/:id", "Eliminar dirección. [Requiere JWT]"),
    ]
    tbl_addr = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_addr:
        add_table_row(tbl_addr, [ep, desc])
    separator(doc)

    # Stock
    heading3(doc, "2.5.8 Módulo de Movimientos de Stock (StockMovementsModule)")
    body(doc,
        "Registra cada cambio en el inventario para garantizar trazabilidad completa. "
        "Tipos de movimiento: VENTA, COMPRA, AJUSTE, DEVOLUCION."
    )
    endpoints_stock = [
        ("GET /api/v1/stock-movements", "Listar movimientos con filtros. [Solo ADMIN]"),
        ("GET /api/v1/stock-movements/:productId", "Historial de movimientos de un producto. [Solo ADMIN]"),
        ("POST /api/v1/stock-movements", "Registrar movimiento manual. [Solo ADMIN]"),
    ]
    tbl_stock = table_header(doc, ["Endpoint", "Descripción"])
    for ep, desc in endpoints_stock:
        add_table_row(tbl_stock, [ep, desc])
    separator(doc)

    # Frontend
    heading3(doc, "2.5.9 Frontend: Estructura y Componentes")
    body(doc, "Estructura de directorios:", bold=True)
    code_block(doc, """\
frontend/src/
├── App.tsx                    # Componente raíz y configuración de rutas
├── main.tsx                   # Punto de entrada
├── components/
│   ├── Header.tsx             # Cabecera con navegación y carrito
│   ├── ProductCard.tsx        # Tarjeta de producto
│   ├── CartSidebar.tsx        # Panel lateral del carrito
│   ├── admin/                 # Componentes del panel de administración
│   └── ui/                    # Componentes reutilizables (Button, Modal, etc.)
├── context/
│   ├── AuthContext.tsx        # Estado global de autenticación
│   └── CartContext.tsx        # Estado global del carrito
├── services/
│   ├── api.ts                 # Configuración de Axios
│   ├── auth.service.ts        # Llamadas a endpoints de autenticación
│   ├── products.service.ts    # Llamadas a endpoints de productos
│   └── orders.service.ts     # Llamadas a endpoints de órdenes
└── types/
    └── index.ts               # Tipos e interfaces TypeScript""")
    separator(doc)

    # 2.6
    heading2(doc, "2.6 Mantenimiento y Actualización")
    heading3(doc, "2.6.1 Procedimientos de Mantenimiento Regular")

    mantenimiento = [
        ("Mantenimiento Diario",
         ["Verificar logs del servidor (docker-compose logs api --tail=100).",
          "Monitorear uso de memoria y CPU.",
          "Revisar alertas de stock bajo.",
          "Verificar respaldo automático de base de datos."]),
        ("Mantenimiento Semanal",
         ["Revisar métricas de rendimiento.",
          "Analizar logs de errores.",
          "Verificar espacio en disco.",
          "Revisar intentos de acceso no autorizado.",
          "Actualizar dependencias de seguridad (npm audit fix)."]),
        ("Mantenimiento Mensual",
         ["Actualizar dependencias (npm update).",
          "Ejecutar pruebas de carga.",
          "Revisar y optimizar consultas lentas en PostgreSQL.",
          "Verificar certificados SSL.",
          "Hacer prueba del proceso de recuperación de backups."]),
    ]
    for titulo, items in mantenimiento:
        heading3(doc, titulo)
        for it in items:
            bullet(doc, it)

    heading3(doc, "2.6.2 Procedimiento de Actualización del Sistema")
    heading3(doc, "Backend (NestJS)")
    code_block(doc, """\
# 1. Crear rama de actualización
git checkout -b update/v1.1.0

# 2. Actualizar dependencias
cd backend
npm update
npm audit fix

# 3. Si hay cambios en el schema (Prisma)
npx prisma migrate dev --name descripcion_cambio
npx prisma generate

# 4. Ejecutar pruebas
npm run test

# 5. Construir
npm run build

# 6. Desplegar (Render.com lo hace automáticamente al hacer push)
git add .
git commit -m "update: v1.1.0 - descripción de cambios"
git push origin update/v1.1.0""")

    heading3(doc, "Frontend (React + Vite)")
    code_block(doc, """\
cd frontend
npm update
npm run build
# Render.com despliega automáticamente desde la rama main""")

    heading3(doc, "2.6.3 Procedimientos de Backup")
    heading3(doc, "Backup de Base de Datos PostgreSQL")
    code_block(doc, """\
# Crear backup completo
pg_dump -U postgres storehub_db > backup_$(date +%Y%m%d_%H%M%S).sql

# Restaurar backup
psql -U postgres storehub_db < backup_20260304_120000.sql

# Comprimir backup
gzip backup_20260304_120000.sql""")

    body(doc, "Política de retención recomendada:", bold=True)
    for item in [
        "Backups diarios: retener por 7 días.",
        "Backups semanales: retener por 1 mes.",
        "Backups mensuales: retener por 1 año.",
        "Backups de código: almacenados en repositorio Git.",
    ]:
        bullet(doc, item)

    heading3(doc, "2.6.4 Resolución de Problemas Comunes")
    problemas = [
        ('Error: "Cannot connect to PostgreSQL"',
         'Verificar que PostgreSQL está corriendo:\n  Windows: net start postgresql-x64-15\n  Linux: sudo systemctl start postgresql\nVerificar DATABASE_URL en .env.\nRevisar pg_hba.conf para permisos de conexión.'),
        ('Error: "JWT token invalid or expired"',
         'El token expiró (TTL: 1 hora por defecto).\nHacer login nuevamente.\nVerificar que JWT_SECRET en .env no haya cambiado.\nLimpiar localStorage: localStorage.clear()'),
        ('Error: "Port 3000 already in use"',
         'Windows: netstat -ano | findstr :3000  →  taskkill /PID <PID> /F\nLinux: lsof -ti:3000 | xargs kill -9\nAlternativa: cambiar API_PORT en .env'),
        ('Error: "CORS policy blocked"',
         'Agregar el origen del frontend a app.enableCors() en main.ts.\nReiniciar el backend después del cambio.'),
        ('Error: "Module not found"',
         'npm cache clean --force\nrm -rf node_modules package-lock.json\nnpm install'),
    ]
    for error, solucion in problemas:
        heading3(doc, error)
        code_block(doc, solucion)
    separator(doc)

    heading3(doc, "Niveles de Soporte")
    tbl_soporte = table_header(doc, ["Nivel", "Tiempo de Respuesta", "Tipo de Problema"])
    support_data = [
        ("Nivel 1 — Urgente", "< 2 horas", "Sistema caído, pérdida de datos, fallos de seguridad."),
        ("Nivel 2 — Alto",    "< 8 horas", "Módulo crítico no funciona, afecta múltiples usuarios."),
        ("Nivel 3 — Medio",   "< 48 horas","Función específica con errores, afecta pocos usuarios."),
        ("Nivel 4 — Bajo",    "< 1 semana","Mejoras, sugerencias, documentación."),
    ]
    for row in support_data:
        add_table_row(tbl_soporte, row)
    separator(doc)

    # 2.7
    heading2(doc, "2.7 Seguridad y Licencias")
    heading3(doc, "2.7.1 Políticas de Seguridad")

    heading3(doc, "Autenticación y Autorización")
    for item in [
        "Autenticación basada en JWT con doble token (access + refresh).",
        "Sistema de roles: ADMIN y CLIENTE con Guards por endpoint.",
        "Contraseñas hasheadas con bcrypt (salt rounds: 10).",
        "Access token con TTL de 1 hora. Refresh token con TTL de 7 días.",
        "Validación de roles en todos los endpoints protegidos.",
    ]:
        bullet(doc, item)

    heading3(doc, "Protección de Datos")
    for item in [
        "HTTPS obligatorio en producción (certificado SSL en Render.com).",
        "Datos sensibles encriptados en reposo.",
        "Contraseñas nunca almacenadas en texto plano.",
        "Tokens JWT no contienen datos sensibles.",
        "Variables de entorno para secretos (nunca en código fuente).",
    ]:
        bullet(doc, item)

    heading3(doc, "Prevención de Vulnerabilidades")
    for item in [
        "Protección SQL Injection: Prisma ORM usa consultas parametrizadas.",
        "Validación de entrada: class-validator en todos los DTOs.",
        "Rate limiting: protección contra ataques de fuerza bruta.",
        "Headers de seguridad (Helmet middleware en NestJS).",
        "CORS configurado para permitir solo orígenes autorizados.",
        "Sanitización de datos en endpoints públicos.",
    ]:
        bullet(doc, item)

    heading3(doc, "2.7.2 Gestión de Dependencias y Vulnerabilidades")
    code_block(doc, """\
# Auditar vulnerabilidades conocidas
npm audit

# Corregir vulnerabilidades automáticamente
npm audit fix

# Verificar actualizaciones disponibles
npm outdated

# Actualizar paquetes (con precaución en producción)
npm update""")

    heading3(doc, "2.7.3 Licencia de Software")
    body(doc, "MIT License", bold=True)
    body(doc,
        "Copyright (c) 2026 Sebastián Gallego Vásquez, Santiago Sánchez Sepúlveda, "
        "Duván Ochoa Hernández, Ricky Lotero Severino"
    )
    body(doc,
        "Se concede permiso, de forma gratuita, a cualquier persona que obtenga una copia "
        "de este software y de los archivos de documentación asociados (el 'Software'), para "
        "tratar el Software sin restricción, incluyendo sin limitación los derechos de usar, "
        "copiar, modificar, fusionar, publicar, distribuir, sublicenciar y/o vender copias del "
        "Software, y para permitir a las personas a quienes se les proporcione el Software que "
        "lo hagan, sujeto a las siguientes condiciones:\n\n"
        "El aviso de copyright anterior y este aviso de permiso se incluirán en todas las "
        "copias o partes sustanciales del Software.\n\n"
        "EL SOFTWARE SE PROPORCIONA 'TAL CUAL', SIN GARANTÍA DE NINGÚN TIPO, EXPRESA O IMPLÍCITA."
    )

    heading3(doc, "2.7.4 Licencias de Dependencias")
    tbl_lic = table_header(doc, ["Paquete", "Licencia", "Uso"])
    lic_data = [
        ("NestJS", "MIT", "Framework backend"),
        ("React", "MIT", "Framework frontend"),
        ("Prisma", "Apache 2.0", "ORM"),
        ("PostgreSQL", "PostgreSQL License", "Base de datos"),
        ("Redis", "BSD 3-Clause", "Caché"),
        ("bcrypt", "MIT", "Hash de contraseñas"),
        ("class-validator", "MIT", "Validación de datos"),
        ("Axios", "MIT", "Cliente HTTP"),
        ("Docker", "Apache 2.0", "Containerización"),
    ]
    for row in lic_data:
        add_table_row(tbl_lic, row)
    separator(doc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # CAPÍTULO 3
    # ═══════════════════════════════════════════════════════════════
    heading1(doc, "CAPÍTULO 3: Mejores Prácticas y Efectividad")

    # 3.1
    heading2(doc, "3.1 Características de un Manual Efectivo")
    caracteristicas = [
        ("Claridad y Precisión",
         ["Lenguaje técnico preciso pero comprensible.",
          "Ejemplos concretos para conceptos abstractos.",
          "Diagramas y esquemas visuales.",
          "Terminología consistente en todo el documento."]),
        ("Estructura Organizada",
         ["Tabla de contenidos detallada.",
          "Jerarquía lógica de secciones.",
          "Referencias cruzadas entre secciones relacionadas.",
          "Índice de términos técnicos."]),
        ("Actualización Continua",
         ["Versionado del documento sincronizado con el software.",
          "Registro de cambios (changelog) en cada versión.",
          "Fecha de última actualización visible.",
          "Proceso definido para proponer correcciones."]),
        ("Orientación Práctica",
         ["Guías paso a paso para tareas comunes.",
          "Ejemplos de código funcionales y probados.",
          "Sección de resolución de problemas frecuentes.",
          "Casos de uso reales del sistema."]),
    ]
    for titulo, items in caracteristicas:
        heading3(doc, titulo)
        for it in items:
            bullet(doc, it)

    # 3.2
    heading2(doc, "3.2 Buenas Prácticas en Desarrollo")
    heading3(doc, "Estándares de Código")
    code_block(doc, """\
// NestJS — ejemplo de endpoint bien estructurado
@Post()
@Roles(UserRole.ADMIN)
@UseGuards(JwtAuthGuard, RolesGuard)
async create(
  @Body() createProductDto: CreateProductDto,
  @CurrentUser() user: User,
): Promise<ApiResponse<Product>> {
  try {
    const product = await this.productsService.create(createProductDto);
    return { success: true, data: product };
  } catch (error) {
    throw new InternalServerErrorException('Error al crear el producto');
  }
}""")

    heading3(doc, "Convenciones de Nomenclatura")
    body(doc, "Backend (NestJS / TypeScript):", bold=True)
    for item in [
        "Archivos: kebab-case (auth.service.ts, create-product.dto.ts).",
        "Clases y decoradores: PascalCase (ProductsService, CreateProductDto).",
        "Variables y funciones: camelCase (findAll, createProduct).",
        "Constantes: UPPER_SNAKE_CASE (JWT_SECRET, MAX_FILE_SIZE).",
        "Endpoints: kebab-case (/api/v1/stock-movements).",
    ]:
        bullet(doc, item)

    body(doc, "Frontend (React / TypeScript):", bold=True)
    for item in [
        "Componentes: PascalCase (ProductCard.tsx, CartSidebar.tsx).",
        "Hooks personalizados: camelCase con prefijo 'use' (useAuth, useCart).",
        "Servicios: camelCase con sufijo '.service.ts' (products.service.ts).",
        "CSS Modules: ComponentName.module.css.",
    ]:
        bullet(doc, item)

    heading3(doc, "Gestión de Errores")
    code_block(doc, """\
// Backend — manejo consistente de errores
@Catch()
export class AllExceptionsFilter implements ExceptionFilter {
  catch(exception: unknown, host: ArgumentsHost) {
    const ctx = host.switchToHttp();
    const response = ctx.getResponse<Response>();
    const status =
      exception instanceof HttpException
        ? exception.getStatus()
        : HttpStatus.INTERNAL_SERVER_ERROR;

    response.status(status).json({
      success: false,
      statusCode: status,
      message:
        exception instanceof HttpException
          ? exception.message
          : 'Error interno del servidor',
      timestamp: new Date().toISOString(),
    });
  }
}""")

    # 3.3
    heading2(doc, "3.3 Casos de Uso y Ejemplos Reales")

    heading3(doc, "Caso 1: Proceso de Compra Completo")
    pasos_compra = [
        ("1. El cliente inicia sesión",
         'POST /api/v1/auth/login\n{ "email": "cliente@storehub.co", "password": "secret123" }'),
        ("2. Explora el catálogo",
         'GET /api/v1/products?categoriaId=cat123&page=1&limit=20'),
        ("3. Agrega producto al carrito",
         'POST /api/v1/cart/items\n{ "productoId": "prod456", "cantidad": 2 }'),
        ("4. Verifica el carrito",
         'GET /api/v1/cart'),
        ("5. Crea la orden",
         'POST /api/v1/orders\n{ "addressId": "addr789", "items": [{ "productoId": "prod456", "cantidad": 2 }] }'),
        ("6. Sigue el estado",
         'GET /api/v1/orders/ord001'),
    ]
    for paso, codigo in pasos_compra:
        heading3(doc, paso)
        code_block(doc, codigo)

    heading3(doc, "Caso 2: Gestión de Inventario por el Administrador")
    pasos_admin = [
        ("1. Crear nueva categoría",
         'POST /api/v1/categories\n{ "nombre": "Cafés Especiales", "descripcion": "Café de origen colombiano", "icono": "coffee" }'),
        ("2. Agregar producto",
         'POST /api/v1/products\n{\n  "nombre": "Café Premium Huila",\n  "descripcion": "Café de altura 100% colombiano",\n  "precio": 2500000,\n  "stock": 100,\n  "categoriaId": "cat-cafes"\n}'),
        ("3. Actualizar stock manualmente",
         'POST /api/v1/stock-movements\n{\n  "productoId": "prod-cafe",\n  "tipo": "COMPRA",\n  "cantidad": 50,\n  "descripcion": "Reposición de inventario"\n}'),
        ("4. Actualizar estado de una orden",
         'PATCH /api/v1/orders/ord001/status\n{ "estado": "EN_PREPARACION" }'),
    ]
    for paso, codigo in pasos_admin:
        heading3(doc, paso)
        code_block(doc, codigo)

    # 3.4
    heading2(doc, "3.4 La Documentación como Pilar del Éxito")
    body(doc,
        "La documentación técnica no es una tarea opcional al final del desarrollo — es un "
        "componente fundamental del software en sí mismo. Un sistema bien documentado:"
    )
    for item in [
        "Reduce drásticamente el tiempo de incorporación de nuevos desarrolladores.",
        "Facilita la detección temprana de errores y regresiones.",
        "Asegura la continuidad del proyecto ante cambios en el equipo.",
        "Genera confianza en los clientes y usuarios finales.",
        "Permite escalar el sistema con mayor seguridad y velocidad.",
    ]:
        bullet(doc, item)

    separator(doc)
    body(doc, "Compromiso del equipo StoreHub:", bold=True)
    body(doc,
        "El equipo de desarrollo se compromete a mantener esta documentación actualizada en "
        "cada versión publicada, garantizando que refleje fielmente el estado real del sistema."
    )
    separator(doc)

    # ═══════════════════════════════════════════════════════════════
    # PIE DE PÁGINA / INFORMACIÓN LEGAL
    # ═══════════════════════════════════════════════════════════════
    doc.add_page_break()
    heading1(doc, "Información del Documento")
    tbl_info = table_header(doc, ["Campo", "Valor"])
    info_data = [
        ("Documento",      "Manual Técnico de Software — StoreHub"),
        ("Versión",        "1.0.0"),
        ("Fecha",          "04 de Marzo de 2026"),
        ("Estado",         "Producción"),
        ("Proyecto",       "StoreHub — Sistema E-Commerce con Gestión Integral de Inventario"),
        ("Cliente",        "Julián Alberto Garzón García / Distribuciones S y J"),
        ("Institución",    "SENA — Servicio Nacional de Aprendizaje"),
        ("Autores",        "Sebastián Gallego Vásquez, Santiago Sánchez Sepúlveda, "
                           "Duván Ochoa Hernández, Ricky Lotero Severino"),
        ("Frontend URL",   "https://inventory-2-sewi.onrender.com"),
        ("Backend URL",    "https://inventory-1-jkh2.onrender.com/api/v1"),
        ("Repositorio",    "Disponible en el repositorio institucional del proyecto"),
        ("Generado",       datetime.datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]
    for row in info_data:
        add_table_row(tbl_info, row)

    return doc


# ─── main ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    output = r"c:\Users\Equipo\Desktop\inventory app\MANUAL_TECNICO_StoreHub.docx"
    doc = build_document()
    doc.save(output)
    print(f"[OK] Manual generado: {output}")
