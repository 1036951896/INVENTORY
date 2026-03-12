#!/usr/bin/env python3
"""
Generador del documento IEEE 830 - ERS StoreHub
Genera un archivo Word (.docx) profesional con toda la información del proyecto
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colores corporativos ────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x0A, 0x1F, 0x5C)   # Títulos H1
AZUL_MEDIO   = RGBColor(0x1A, 0x56, 0xAB)   # Títulos H2
AZUL_CLARO   = RGBColor(0x2E, 0x7D, 0xC7)   # Títulos H3
GRIS_TEXTO   = RGBColor(0x2C, 0x2C, 0x2C)   # Texto normal
GRIS_TABLA   = RGBColor(0xE8, 0xF0, 0xFB)   # Fondo cabecera tabla
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Pone color de fondo a una celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Agrega bordes finos a toda la tabla"""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{border}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), '1A56AB')
                tcBorders.append(b)
            tcPr.append(tcBorders)

def add_page_break(doc):
    doc.add_page_break()

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_OSCURO
    # línea decorativa bajo el título
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A56AB')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = AZUL_MEDIO

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL_CLARO

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS_TEXTO
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS_TEXTO

def bold_label(doc, label, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = AZUL_OSCURO
    r2 = p.add_run(content)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRIS_TEXTO

def table_header_row(table, headers, widths=None):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, '1A56AB')
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLANCO
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def req_block(doc, numero, nombre, prioridad, descripcion_items):
    """Genera un bloque de requisito funcional con formato de tabla"""
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    set_cell_borders(table)

    # Fila 1: Numero + Nombre
    c = table.cell(0, 0)
    set_cell_bg(c, 'E8F0FB')
    c.paragraphs[0].clear()
    r = c.paragraphs[0].add_run("Numero de requisito")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL_OSCURO

    c2 = table.cell(0, 1)
    c2.paragraphs[0].clear()
    r2 = c2.paragraphs[0].add_run(numero)
    r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = AZUL_MEDIO

    # Fila 2
    c3 = table.cell(1, 0)
    set_cell_bg(c3, 'E8F0FB')
    c3.paragraphs[0].clear()
    r3 = c3.paragraphs[0].add_run("Nombre de requisito")
    r3.bold = True; r3.font.size = Pt(9); r3.font.color.rgb = AZUL_OSCURO

    c4 = table.cell(1, 1)
    c4.paragraphs[0].clear()
    c4.paragraphs[0].add_run(nombre).font.size = Pt(9)

    # Fila 3: Prioridad
    c5 = table.cell(2, 0)
    set_cell_bg(c5, 'E8F0FB')
    c5.paragraphs[0].clear()
    r5 = c5.paragraphs[0].add_run("Prioridad")
    r5.bold = True; r5.font.size = Pt(9); r5.font.color.rgb = AZUL_OSCURO

    c6 = table.cell(2, 1)
    c6.paragraphs[0].clear()
    c6.paragraphs[0].add_run(prioridad).font.size = Pt(9)

    # Fila 4: Descripcion
    c7 = table.cell(3, 0)
    set_cell_bg(c7, 'E8F0FB')
    c7.paragraphs[0].clear()
    r7 = c7.paragraphs[0].add_run("Descripcion")
    r7.bold = True; r7.font.size = Pt(9); r7.font.color.rgb = AZUL_OSCURO

    c8 = table.cell(3, 1)
    c8.paragraphs[0].clear()
    for item in descripcion_items:
        p = c8.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        if item.startswith("- "):
            p.paragraph_format.left_indent = Cm(0.4)
            p.add_run(u"\u2022 " + item[2:]).font.size = Pt(9)
        else:
            p.add_run(item).font.size = Pt(9)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Fuente base
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── PORTADA ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ESPECIFICACIÓN DE REQUISITOS DE SOFTWARE")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = AZUL_OSCURO

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("StoreHub — Plataforma de Comercio Electrónico")
    r2.font.size = Pt(14)
    r2.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph()

    # Tabla portada
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = 'Table Grid'
    set_cell_borders(tbl)
    datos_portada = [
        ("Documento",       "ERS-StoreHub-v1.0"),
        ("Version",         "1.0"),
        ("Fecha",           "Marzo 2026"),
        ("Estandar",        "IEEE 830-1998 (ANSI/IEEE Std 830-1998)"),
        ("Proyecto",        "StoreHub - Plataforma E-Commerce"),
        ("Institucion",     "SENA - Centro de Comercio y Servicios"),
    ]
    for i, (label, val) in enumerate(datos_portada):
        c1 = tbl.cell(i, 0)
        c2 = tbl.cell(i, 1)
        set_cell_bg(c1, '1A56AB')
        c1.paragraphs[0].clear()
        r1 = c1.paragraphs[0].add_run(label)
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = BLANCO
        c2.paragraphs[0].clear()
        c2.paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph()

    # Autores portada
    p_aut = doc.add_paragraph()
    p_aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_aut = p_aut.add_run("Autores: Sebastián Gallego Vásquez  |  Santiago Sánchez Sepúlveda  |  Duván Ochoa Hernández  |  Ricky Lotero Severino")
    r_aut.font.size = Pt(10)
    r_aut.italic = True
    r_aut.font.color.rgb = GRIS_TEXTO

    add_page_break(doc)

    # ── FICHA DEL DOCUMENTO ────────────────────────────────────────────────────
    heading1(doc, "Ficha del Documento")

    tbl2 = doc.add_table(rows=3, cols=4)
    tbl2.style = 'Table Grid'
    set_cell_borders(tbl2)
    table_header_row(tbl2, ["Fecha", "Revision", "Autores", "Verificado"])

    filas_ficha = [
        ("07/2025", "1 — Inicial",
         "Sebastián Gallego\nSantiago Sánchez\nDuván Ochoa\nRicky Lotero",
         "En revisión"),
        ("03/2026", "2 — Final",
         "Sebastián Gallego\nSantiago Sánchez\nDuván Ochoa\nRicky Lotero",
         "Aprobado"),
    ]
    for i, (fecha, rev, autores, ver) in enumerate(filas_ficha):
        row = tbl2.rows[i + 1]
        for j, val in enumerate([fecha, rev, autores, ver]):
            row.cells[j].paragraphs[0].clear()
            row.cells[j].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph()
    body(doc, "Documento validado por las partes el: Marzo 2026")

    tbl3 = doc.add_table(rows=2, cols=2)
    tbl3.style = 'Table Grid'
    set_cell_borders(tbl3)
    table_header_row(tbl3, ["Por el Cliente", "Por la Institucion"])
    tbl3.cell(1, 0).paragraphs[0].clear()
    tbl3.cell(1, 0).paragraphs[0].add_run(
        "Julián Alberto Garzón García\nDistribuciones S y J"
    ).font.size = Pt(9)
    tbl3.cell(1, 1).paragraphs[0].clear()
    tbl3.cell(1, 1).paragraphs[0].add_run(
        "SENA — Centro de Comercio y Servicios\nInstructor: [Nombre del instructor]"
    ).font.size = Pt(9)

    add_page_break(doc)

    # ── SECCIÓN 1: INTRODUCCIÓN ────────────────────────────────────────────────
    heading1(doc, "1. Introducción")

    body(doc,
        "Este documento corresponde a la Especificación de Requisitos de Software (ERS) "
        "para el sistema StoreHub, una plataforma de comercio electrónico desarrollada "
        "como proyecto formativo en el SENA. Esta especificación ha sido elaborada "
        "conforme a las directrices del estándar IEEE 830 (ANSI/IEEE Std 830-1998), "
        "Práctica Recomendada para la Especificación de Requisitos de Software.")

    # 1.1 Propósito
    heading2(doc, "1.1 Propósito")
    body(doc,
        "Desarrollar StoreHub, una aplicación web de comercio electrónico orientada a "
        "pequeñas y medianas empresas, que permita la gestión integral de productos, "
        "inventario, pedidos y clientes, con herramientas de visualización y reportes "
        "para la toma de decisiones estratégicas.")
    body(doc, "El presente documento va dirigido a:")
    bullet(doc, "El equipo de desarrollo del proyecto (programadores, analistas, diseñadores).")
    bullet(doc, "El cliente: Julián Alberto Garzón García / Distribuciones S y J.")
    bullet(doc, "Evaluadores académicos del SENA (Centro de Comercio y Servicios).")

    # 1.2 Alcance
    heading2(doc, "1.2 Alcance")
    body(doc,
        "StoreHub es una aplicación web de comercio electrónico orientada a la gestión "
        "eficiente de inventarios, ventas y clientes. Está diseñada para adaptarse a "
        "distintos tipos de negocios, siendo Distribuciones S y J el cliente principal "
        "del sistema.")
    body(doc, "Las funcionalidades principales son:")
    bullet(doc, "Tienda en línea con catálogo de productos, ofertas y búsqueda por categorías.")
    bullet(doc, "Gestión de inventario: registro, edición, eliminación y alertas de stock bajo.")
    bullet(doc, "Carrito de compras y flujo completo de pedidos con confirmación.")
    bullet(doc, "Panel administrativo con dashboard, métricas e indicadores en tiempo real.")
    bullet(doc, "Control de acceso por roles: administrador, cliente e invitado.")
    bullet(doc, "Módulo de reportes y exportación de datos.")
    bullet(doc, "Gestión de ofertas y promociones.")
    body(doc,
        "El sistema busca optimizar el control de existencias, mejorar la trazabilidad de "
        "productos y facilitar los procesos de compra, venta y gestión de pedidos para "
        "negocios locales y en expansión.")

    # 1.3 Personal involucrado
    heading2(doc, "1.3 Personal Involucrado")

    # Tabla resumen
    tbl4 = doc.add_table(rows=6, cols=4)
    tbl4.style = 'Table Grid'
    set_cell_borders(tbl4)
    table_header_row(tbl4, ["Nombre Completo", "Rol", "Correo", "Telefono"])
    personal = [
        ("Julián Alberto Garzón García", "Cliente / Distribuciones S y J", "—", "3122808925"),
        ("Santiago Sánchez Sepúlveda",   "Tecnólogo - Analista y Programador", "santis3268@gmail.com", "3116579677"),
        ("Sebastián Gallego Vásquez",    "Tecnólogo - Analista y Programador", "johnsesistemas@gmail.com", "3133856206"),
        ("Duván Ochoa Hernández",        "Tecnólogo - Analista y Programador", "duvan_aochoa@soy.sena.edu.co", "3194171074"),
        ("Ricky Lotero Severino",        "Tecnólogo - Analista y Programador", "rikylotero@gmail.com", "3135468253"),
    ]
    for i, row_data in enumerate(personal):
        row = tbl4.rows[i + 1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        for j, val in enumerate(row_data):
            row.cells[j].paragraphs[0].clear()
            row.cells[j].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph()

    # Fichas individuales
    fichas = [
        {
            "nombre": "Sebastián Gallego Vásquez",
            "rol": "Analista, diseñador y programador",
            "categoria": "Tecnólogo en Análisis y Desarrollo de Software — SENA",
            "responsabilidades": "Análisis de requerimientos, diseño de arquitectura frontend, "
                                 "desarrollo de componentes React, integración con APIs, implementación "
                                 "del sistema de autenticación y carrito de compras en StoreHub.",
            "contacto": "johnsesistemas@gmail.com | 3133856206",
        },
        {
            "nombre": "Santiago Sánchez Sepúlveda",
            "rol": "Analista, diseñador y programador",
            "categoria": "Tecnólogo en Análisis y Desarrollo de Software — SENA",
            "responsabilidades": "Análisis de información, diseño de experiencia de usuario (UX/UI), "
                                 "desarrollo de módulos de gestión de productos, panel administrativo "
                                 "y sistema de reportes en StoreHub.",
            "contacto": "santis3268@gmail.com | 3116579677",
        },
        {
            "nombre": "Duván Ochoa Hernández",
            "rol": "Analista, diseñador y programador",
            "categoria": "Tecnólogo en Análisis y Desarrollo de Software — SENA",
            "responsabilidades": "Análisis de información, diseño y desarrollo del backend con NestJS, "
                                 "modelado de base de datos PostgreSQL con Prisma ORM, implementación "
                                 "de endpoints REST y despliegue en la nube para StoreHub.",
            "contacto": "duvan_aochoa@soy.sena.edu.co | 3194171074",
        },
        {
            "nombre": "Ricky Lotero Severino",
            "rol": "Analista, diseñador y programador",
            "categoria": "Tecnólogo en Análisis y Desarrollo de Software — SENA",
            "responsabilidades": "Análisis de información, diseño visual, desarrollo del módulo de "
                                 "pedidos, seguimiento de órdenes, pruebas funcionales e integración "
                                 "del sistema completo de StoreHub.",
            "contacto": "rikylotero@gmail.com | 3135468253",
        },
    ]

    for ficha in fichas:
        doc.add_paragraph()
        tbl_f = doc.add_table(rows=5, cols=2)
        tbl_f.style = 'Table Grid'
        set_cell_borders(tbl_f)
        campos = [
            ("Nombre", ficha["nombre"]),
            ("Rol", ficha["rol"]),
            ("Categoria profesional", ficha["categoria"]),
            ("Responsabilidades", ficha["responsabilidades"]),
            ("Informacion de contacto", ficha["contacto"]),
        ]
        for i, (lbl, val) in enumerate(campos):
            c1 = tbl_f.cell(i, 0)
            c2 = tbl_f.cell(i, 1)
            set_cell_bg(c1, 'E8F0FB')
            c1.paragraphs[0].clear()
            r1 = c1.paragraphs[0].add_run(lbl)
            r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = AZUL_OSCURO
            c2.paragraphs[0].clear()
            c2.paragraphs[0].add_run(val).font.size = Pt(9)

    # 1.4 Definiciones y acrónimos
    heading2(doc, "1.4 Definiciones, Acrónimos y Abreviaturas")
    tbl5 = doc.add_table(rows=14, cols=2)
    tbl5.style = 'Table Grid'
    set_cell_borders(tbl5)
    table_header_row(tbl5, ["Termino / Acronimo", "Definicion"])
    terminos = [
        ("StoreHub",     "Plataforma web de comercio electrónico para gestión de inventarios, productos y pedidos."),
        ("ERS",          "Especificación de Requisitos de Software."),
        ("IEEE 830",     "Estándar para la especificación de requisitos de software (ANSI/IEEE Std 830-1998)."),
        ("CRUD",         "Operaciones básicas en bases de datos: Crear, Leer, Actualizar, Eliminar."),
        ("API REST",     "Interfaz de programación de aplicaciones basada en arquitectura REST."),
        ("JWT",          "JSON Web Token — Mecanismo de autenticación segura basado en tokens firmados."),
        ("SPA",          "Single Page Application — Aplicación web de una sola página."),
        ("ORM",          "Object-Relational Mapping — Mapeo objeto-relacional (Prisma en este proyecto)."),
        ("UI / UX",      "Interfaz de Usuario / Experiencia de Usuario."),
        ("Dashboard",    "Panel de control con métricas, gráficas y alertas del sistema."),
        ("Administrador","Usuario con acceso completo a todos los módulos del sistema."),
        ("Cliente",      "Usuario registrado que puede realizar compras y consultar sus pedidos."),
        ("Invitado",     "Usuario sin cuenta que puede navegar el catálogo sin realizar compras."),
    ]
    for i, (term, defn) in enumerate(terminos):
        row = tbl5.rows[i + 1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        row.cells[0].paragraphs[0].clear()
        r_t = row.cells[0].paragraphs[0].add_run(term)
        r_t.bold = True; r_t.font.size = Pt(9)
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run(defn).font.size = Pt(9)

    # 1.5 Referencias
    heading2(doc, "1.5 Referencias")
    refs = [
        ("IEEE Std 830-1998",  "IEEE Recommended Practice for Software Requirements Specifications. IEEE Computer Society, 1998."),
        ("React 19",           "Meta Platforms. React — Biblioteca JavaScript para construcción de interfaces de usuario. https://react.dev"),
        ("NestJS",             "Kamil Mysliwiec. NestJS Framework — Framework Node.js progresivo. https://nestjs.com"),
        ("Prisma ORM",         "Prisma Data Inc. Prisma — Next-generation ORM para Node.js. https://www.prisma.io"),
        ("PostgreSQL 14+",     "The PostgreSQL Global Development Group. Sistema de gestión de base de datos relacional. https://www.postgresql.org"),
        ("Render.com",         "Render Inc. Plataforma de despliegue en la nube. https://render.com"),
        ("Docker",             "Docker Inc. Plataforma de contenedores para despliegue y desarrollo. https://www.docker.com"),
    ]
    for ref, desc in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(ref + ": ")
        r1.bold = True; r1.font.size = Pt(9)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9); r2.font.color.rgb = GRIS_TEXTO

    # 1.6 Resumen
    heading2(doc, "1.6 Resumen")
    body(doc,
        "Este documento detalla los requisitos del sistema StoreHub, una solución tecnológica "
        "moderna para gestión de inventario, ventas en línea y administración de clientes, "
        "desarrollada con tecnologías de vanguardia (React, NestJS, PostgreSQL) para "
        "Distribuciones S y J y negocios similares. Se describen los objetivos, "
        "funcionalidades, características de usuarios, restricciones técnicas y la "
        "evolución prevista del sistema. Su estructura sigue el estándar IEEE 830 para "
        "facilitar la trazabilidad y comprensión por todas las partes involucradas.")
    body(doc,
        "El documento está organizado en tres grandes bloques: introducción (sección 1), "
        "descripción general del sistema (sección 2) y requisitos específicos funcionales "
        "y no funcionales (sección 3).")

    add_page_break(doc)

    # ── SECCIÓN 2: DESCRIPCIÓN GENERAL ────────────────────────────────────────
    heading1(doc, "2. Descripción General")

    # 2.1 Perspectiva del producto
    heading2(doc, "2.1 Perspectiva del Producto")
    body(doc,
        "StoreHub es una solución de software independiente de comercio electrónico que "
        "puede operar de forma autónoma en entornos comerciales de pequeña y mediana "
        "escala. Está construida sobre una arquitectura de tres capas con separación "
        "clara de responsabilidades:")
    bullet(doc, "Frontend (Capa de Presentación): Aplicación React 19 + TypeScript (SPA), "
                "desplegada en Render.com como servidor Express. Accesible desde cualquier "
                "navegador moderno con diseño completamente responsivo.")
    bullet(doc, "Backend (Capa de Negocio): API REST desarrollada con NestJS + TypeScript, "
                "con autenticación JWT, validación de datos, control de roles y lógica de "
                "negocio centralizada. Desplegada en Render.com con Docker.")
    bullet(doc, "Base de Datos (Capa de Datos): PostgreSQL 14+ gestionada con Prisma ORM. "
                "Alojada en Render.com como base de datos administrada.")
    body(doc,
        "El sistema está diseñado con enfoque modular para permitir escalar de forma "
        "controlada, y ya cuenta con despliegue en producción en la nube (Render.com), "
        "lo que permite acceso 24/7 desde cualquier dispositivo con conexión a internet.")

    # 2.2 Funcionalidad
    heading2(doc, "2.2 Funcionalidad del Producto")
    body(doc, "StoreHub ofrece las siguientes funcionalidades principales:")
    funciones = [
        "Tienda en linea: catálogo de productos con filtros por categoria, busqueda, detalle de producto y visita de invitados sin registro.",
        "Gestion de inventario: registro, edicion, eliminacion y control de productos con alertas automaticas de stock bajo.",
        "Carrito de compras: agregar/eliminar productos, modificar cantidades, calcular totales y persistencia durante la sesion.",
        "Flujo de pedidos: confirmacion, resumen de compra, reduccion automatica de inventario y generacion de comprobante.",
        "Seguimiento de pedidos: historial del cliente con estado de cada orden.",
        "Panel administrativo (Dashboard): visualizacion de metricas clave, graficas de ventas, productos mas vendidos y alertas.",
        "Gestion de usuarios: registro, autenticacion JWT, control de roles (admin / cliente / invitado) y perfiles.",
        "Modulo de ofertas: creacion y gestion de promociones y descuentos para productos seleccionados.",
        "Reportes: visualizacion de datos de ventas, inventario y pedidos con filtros.",
        "Gestion de categorias: creacion y administracion de categorias de productos.",
    ]
    for f in funciones:
        bullet(doc, f)

    # 2.3 Características de los usuarios
    heading2(doc, "2.3 Caracteristicas de los Usuarios")

    tbl6 = doc.add_table(rows=4, cols=4)
    tbl6.style = 'Table Grid'
    set_cell_borders(tbl6)
    table_header_row(tbl6, ["Tipo de Usuario", "Formacion", "Habilidades", "Actividades"])
    usuarios = [
        ("Administrador",
         "Conocimientos basicos o medios en informatica y gestion de inventarios.",
         "Capacidad para registrar, editar y eliminar productos, gestionar pedidos, "
         "usuarios, categorias y generar reportes.",
         "Gestiona el inventario, supervisa compras y pedidos, controla usuarios, "
         "consulta el dashboard y exporta reportes."),
        ("Cliente",
         "Usuario general sin necesidad de formacion tecnica.",
         "Navegacion web, compras en linea basicas.",
         "Navega el catalogo, agrega al carrito, realiza pedidos y consulta "
         "su historial de ordenes."),
        ("Invitado",
         "Sin requisitos de formacion.",
         "Uso basico de interfaces web.",
         "Consulta el catalogo y ofertas publicas sin realizar compras "
         "ni acceder a informacion sensible."),
    ]
    for i, row_data in enumerate(usuarios):
        row = tbl6.rows[i + 1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        for j, val in enumerate(row_data):
            row.cells[j].paragraphs[0].clear()
            row.cells[j].paragraphs[0].add_run(val).font.size = Pt(9)

    # 2.4 Restricciones
    heading2(doc, "2.4 Restricciones")

    heading3(doc, "Lenguajes y Tecnologias")
    restricciones_tec = [
        "El frontend se desarrolla con React 19 + TypeScript compilado con Vite.",
        "El backend se desarrolla con NestJS (Node.js) + TypeScript.",
        "La base de datos es PostgreSQL 14+ administrada mediante Prisma ORM.",
        "La autenticacion se gestiona exclusivamente con JSON Web Tokens (JWT).",
        "El despliegue se realiza en Render.com usando contenedores Docker para el backend.",
        "Los estilos del frontend usan CSS modular propio con sistema de diseno personalizado.",
    ]
    for r in restricciones_tec:
        bullet(doc, r)

    heading3(doc, "Restricciones de Desarrollo")
    restricciones_dev = [
        "Se sigue un enfoque iterativo e incremental sin aplicar metodologias agiles formales, dado el contexto formativo del proyecto.",
        "El equipo esta conformado por aprendices del SENA, por lo tanto el avance esta sujeto a la disponibilidad de tiempo y objetivos de formacion.",
        "No se contempla integracion con sistemas de facturacion electronica ni pasarelas de pago reales en esta version.",
    ]
    for r in restricciones_dev:
        bullet(doc, r)

    heading3(doc, "Restricciones de Despliegue")
    restricciones_desp = [
        "El sistema requiere un entorno con Node.js 18+ para ejecutar el backend.",
        "La base de datos debe ser PostgreSQL 14 o superior.",
        "Para desarrollo local se utiliza Docker Compose con los servicios backend y base de datos.",
        "No se contempla desarrollo de aplicacion movil nativa en esta version; el enfoque es exclusivamente web responsivo.",
    ]
    for r in restricciones_desp:
        bullet(doc, r)

    # 2.5 Suposiciones y dependencias
    heading2(doc, "2.5 Suposiciones y Dependencias")

    heading3(doc, "Suposiciones")
    suposiciones = [
        "Se asume que los usuarios (administradores y clientes) cuentan con conocimientos basicos en el uso de plataformas web y navegadores modernos.",
        "Se espera que la empresa usuaria (Distribuciones S y J) proporcione la informacion necesaria: productos, precios, stock inicial y datos de contacto.",
        "El entorno de produccion contara con acceso estable a internet para el correcto funcionamiento del sistema en la nube.",
        "Los usuarios acceden al sistema desde navegadores modernos: Google Chrome, Mozilla Firefox, Microsoft Edge o Safari (ultimas 2 versiones).",
        "Las imagenes de productos seran provistas por el administrador en formatos web estandar (JPG, PNG, WebP).",
    ]
    for s in suposiciones:
        bullet(doc, s)

    heading3(doc, "Dependencias")
    dependencias = [
        "El frontend depende de la API REST del backend para todas las operaciones de datos.",
        "El backend depende de Prisma ORM y de una instancia activa de PostgreSQL.",
        "El sistema de autenticacion depende de la correcta configuracion de variables de entorno (JWT_SECRET, DATABASE_URL).",
        "El despliegue en produccion depende de la disponibilidad del servicio Render.com.",
        "Se espera acceso a internet de minimo 5 Mbps para un uso fluido del sistema.",
        "El uso del panel administrativo requiere que el usuario tenga rol 'admin' correctamente asignado en la base de datos.",
    ]
    for d in dependencias:
        bullet(doc, d)

    # 2.6 Evolución previsible
    heading2(doc, "2.6 Evolucion Previsible del Sistema")
    body(doc,
        "Se espera que StoreHub evolucione de forma progresiva conforme a las necesidades "
        "del cliente y la madurez del proyecto. Las posibles mejoras y ampliaciones futuras son:")
    evoluciones = [
        "Integracion con pasarelas de pago (PSE, Nequi, tarjeta de credito) para habilitar cobros en linea.",
        "Soporte multitienda (multi-tenant), permitiendo que diferentes empresas usen la misma plataforma con datos independientes.",
        "Modulo de auditoria que registre todas las operaciones realizadas por los usuarios del sistema.",
        "Notificaciones por correo electronico (confirmacion de pedidos, alertas de stock, cambios de estado).",
        "Exportacion avanzada de reportes en PDF y Excel con graficas incluidas.",
        "Aplicacion movil (PWA o nativa) como complemento a la version web.",
        "Integracion con sistemas de contabilidad y facturacion electronica segun normativa colombiana.",
        "Motor de busqueda avanzado con filtros multiples, busqueda por voz y sugerencias.",
        "Sistema de valoraciones y resenas de productos por parte de los clientes.",
        "Optimizacion y escalabilidad con migracion a microservicios si el volumen de operaciones lo requiere.",
    ]
    for e in evoluciones:
        bullet(doc, e)

    add_page_break(doc)

    # ── SECCIÓN 3: REQUISITOS ESPECÍFICOS ─────────────────────────────────────
    heading1(doc, "3. Requisitos Específicos")

    # 3.1 Interfaces comunes
    heading2(doc, "3.1 Requisitos Comunes de los Interfaces")

    heading3(doc, "Entradas del Sistema")
    body(doc, "Las principales entradas que los usuarios proporcionan al sistema son:")
    entradas = [
        "Formulario de inicio de sesion: correo electronico y contrasena.",
        "Formulario de registro: nombre completo, correo, contrasena y telefono.",
        "Formulario de producto (admin): nombre, descripcion, precio, cantidad en stock, categoria, imagenes y estado (activo/inactivo).",
        "Seleccion de productos en tienda: cantidad deseada y boton 'Agregar al carrito'.",
        "Formulario de confirmacion de pedido: datos de contacto y confirmacion del total.",
        "Filtros de busqueda: texto de busqueda, categoria y rango de precio.",
        "Formulario de oferta (admin): nombre, descripcion, porcentaje de descuento y productos aplicables.",
        "Formulario de categoria (admin): nombre e imagen de categoria.",
    ]
    for e in entradas:
        bullet(doc, e)

    heading3(doc, "Salidas del Sistema")
    body(doc, "Las respuestas y elementos generados por el sistema son:")
    salidas = [
        "Mensajes de autenticacion: exito o error al iniciar/cerrar sesion.",
        "Catalogo de productos: tarjetas con imagen, nombre, precio, categoria y boton de accion.",
        "Vista de carrito: productos seleccionados, cantidades editables, subtotales y total.",
        "Resumen de pedido confirmado: productos, cantidades, total, fecha y numero de orden.",
        "Dashboard administrativo: estadisticas en tarjetas, graficas de ventas y alertas de stock bajo.",
        "Historial de pedidos del cliente: lista de ordenes con estado y detalle.",
        "Reportes de inventario: tabla de productos con stock actual, entradas y salidas.",
        "Alertas de stock bajo: lista de productos cuya cantidad esta por debajo del umbral configurado.",
        "Notificaciones en pantalla (toast): confirmaciones, errores y advertencias en tiempo real.",
    ]
    for s in salidas:
        bullet(doc, s)

    # 3.1.1 Interfaces de usuario
    heading3(doc, "3.1.1 Interfaces de Usuario")
    body(doc,
        "La interfaz de StoreHub ha sido diseñada para ser clara, intuitiva y completamente "
        "responsiva, adaptandose a escritorio, tableta y movil. Esta orientada a tres "
        "perfiles: administrador, cliente e invitado.")
    body(doc, "Diseno general:")
    bullet(doc, "Paleta de colores moderna con azul profundo (#0A1F5C) como color primario, "
                "blancos y grises para fondos y naranja/dorado para llamadas a la accion.")
    bullet(doc, "Tipografia sans-serif moderna (Inter, Roboto) con jerarquia visual clara.")
    bullet(doc, "Componentes React reutilizables con CSS modular y sistema de diseno personalizado.")
    bullet(doc, "Diseño mobile-first con breakpoints para tablets y escritorio.")

    body(doc, "Pantallas clave del sistema:")
    pantallas = [
        "Pagina de inicio / Landing: hero section, categorias destacadas, productos en oferta y catalogo general.",
        "Catalogo de productos: grilla de tarjetas con filtros por categoria, busqueda en tiempo real y paginacion.",
        "Detalle de producto: imagen, descripcion completa, precio, stock disponible y boton de compra.",
        "Carrito de compras (panel lateral): lista de items, cantidades editables, total y boton de checkout.",
        "Checkout / Confirmacion de pedido: resumen de compra, datos de contacto y confirmacion.",
        "Seguimiento de pedidos: historial de ordenes del cliente con estado actualizable.",
        "Login / Registro: formularios limpios con validacion en tiempo real y mensajes de error.",
        "Panel Admin - Dashboard: metricas en tarjetas (ventas, pedidos, productos, usuarios), graficas.",
        "Panel Admin - Productos: tabla con CRUD completo, busqueda, filtros y paginacion.",
        "Panel Admin - Pedidos: lista de ordenes con cambio de estado y detalles.",
        "Panel Admin - Categorias: gestion de categorias con imagen.",
        "Panel Admin - Usuarios: lista de usuarios registrados con gestion de roles.",
        "Panel Admin - Ofertas: creacion y administracion de promociones y descuentos.",
        "Pagina de Ofertas Publicas: vitrina de productos en promocion accesible sin login.",
    ]
    for p_item in pantallas:
        bullet(doc, p_item)

    # 3.1.2 Interfaces de hardware
    heading3(doc, "3.1.2 Interfaces de Hardware")
    body(doc, "Dado que StoreHub es una aplicacion web, las dependencias de hardware son minimas.")

    body(doc, "Lado cliente (usuarios: administrador, cliente, invitado):")
    tbl7 = doc.add_table(rows=7, cols=2)
    tbl7.style = 'Table Grid'
    set_cell_borders(tbl7)
    table_header_row(tbl7, ["Elemento", "Descripcion minima requerida"])
    hw_cliente = [
        ("Dispositivo",       "PC, portatil, tableta o smartphone"),
        ("Procesador",        "Dual-core (Intel i3 o equivalente en adelante)"),
        ("Memoria RAM",       "2 GB minimo (recomendado 4 GB)"),
        ("Almacenamiento",    "Solo requiere navegador instalado — sin instalacion local"),
        ("Perifericos",       "Mouse o pantalla tactil, teclado"),
        ("Conectividad",      "Acceso a internet estable (minimo 5 Mbps recomendado)"),
    ]
    for i, (elem, desc) in enumerate(hw_cliente):
        row = tbl7.rows[i + 1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        row.cells[0].paragraphs[0].clear()
        r_e = row.cells[0].paragraphs[0].add_run(elem)
        r_e.bold = True; r_e.font.size = Pt(9)
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(9)

    doc.add_paragraph()
    body(doc, "Lado servidor (produccion en Render.com):")
    tbl8 = doc.add_table(rows=8, cols=2)
    tbl8.style = 'Table Grid'
    set_cell_borders(tbl8)
    table_header_row(tbl8, ["Elemento", "Especificacion"])
    hw_server = [
        ("Sistema operativo",  "Linux (Ubuntu 22.04 en Render.com)"),
        ("Servidor web",       "Node.js 18+ con NestJS y Express"),
        ("Lenguaje backend",   "TypeScript / Node.js 18+"),
        ("Base de datos",      "PostgreSQL 14+ (Managed Database en Render)"),
        ("Contenedor",         "Docker — imagen oficial Node:18-alpine"),
        ("vCPU",               "Minimo 1 vCPU (recomendado 2 vCPU en produccion)"),
        ("RAM servidor",       "Minimo 512 MB (recomendado 1 GB en produccion)"),
    ]
    for i, (elem, desc) in enumerate(hw_server):
        row = tbl8.rows[i + 1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        row.cells[0].paragraphs[0].clear()
        r_e2 = row.cells[0].paragraphs[0].add_run(elem)
        r_e2.bold = True; r_e2.font.size = Pt(9)
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(9)

    # 3.1.3 Interfaces de software
    heading3(doc, "3.1.3 Interfaces de Software")
    body(doc, "StoreHub integra los siguientes productos y servicios de software:")
    sw_interfaces = [
        ("React 19 + Vite",     "Framework frontend para construccion de la SPA. Compilacion y empaquetado de la aplicacion cliente.", "Renderizado y logica de interfaz de usuario."),
        ("NestJS (Node.js)",    "Framework backend modular para la API REST. Maneja rutas, controladores, servicios y validaciones.", "Logica de negocio, autenticacion y acceso a datos."),
        ("Prisma ORM",          "ORM de siguiente generacion para Node.js/TypeScript. Gestiona el esquema, migraciones y consultas.", "Abstraccion de la capa de base de datos sobre PostgreSQL."),
        ("PostgreSQL 14+",      "Sistema gestor de base de datos relacional de codigo abierto.", "Persistencia de todos los datos del sistema."),
        ("JWT (jsonwebtoken)",  "Libreria para generacion y validacion de JSON Web Tokens.", "Autenticacion y autorizacion de usuarios."),
        ("Docker",              "Plataforma de contenedores para empaquetar y desplegar la aplicacion.", "Despliegue consistente en produccion (Render.com)."),
        ("Render.com",          "Plataforma cloud PaaS para despliegue del frontend, backend y base de datos.", "Infraestructura de produccion accesible desde internet."),
        ("Sonner (Toasts)",     "Libreria React para notificaciones tipo toast.", "Retroalimentacion visual al usuario en tiempo real."),
    ]
    tbl9 = doc.add_table(rows=len(sw_interfaces)+1, cols=3)
    tbl9.style = 'Table Grid'
    set_cell_borders(tbl9)
    table_header_row(tbl9, ["Producto", "Descripcion", "Proposito"])
    for i, (prod, desc, prop) in enumerate(sw_interfaces):
        row = tbl9.rows[i+1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        row.cells[0].paragraphs[0].clear()
        r_sw = row.cells[0].paragraphs[0].add_run(prod)
        r_sw.bold = True; r_sw.font.size = Pt(9)
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(9)
        row.cells[2].paragraphs[0].clear()
        row.cells[2].paragraphs[0].add_run(prop).font.size = Pt(9)

    # 3.1.4 Interfaces de comunicación
    heading3(doc, "3.1.4 Interfaces de Comunicacion")
    body(doc, "StoreHub utiliza los siguientes protocolos y mecanismos de comunicacion:")
    comms = [
        "HTTP/HTTPS: toda la comunicacion entre el frontend (SPA) y el backend (API REST) se realiza mediante peticiones HTTP con protocolo HTTPS en produccion.",
        "REST API: el backend expone endpoints REST bajo el prefijo /api/v1/ que el frontend consume de forma asincrona con Fetch API.",
        "JSON: formato de intercambio de datos para todas las peticiones y respuestas entre cliente y servidor.",
        "JWT Bearer Token: el token de autenticacion se envia en el header Authorization: Bearer <token> en todas las peticiones protegidas.",
        "CORS configurado: el backend permite solicitudes de origen cruzado (Cross-Origin) desde el dominio del frontend en produccion.",
        "PostgreSQL TCP: comunicacion interna entre el backend (Prisma) y la base de datos PostgreSQL a traves del puerto estandar 5432.",
    ]
    for c in comms:
        bullet(doc, c)

    add_page_break(doc)

    # 3.2 Requisitos funcionales
    heading2(doc, "3.2 Requisitos Funcionales")

    req_block(doc,
        "RF-01", "Inicio de sesion (Login)",
        "Alta / Esencial",
        [
            "El sistema debe permitir que los usuarios registrados (administrador o cliente) "
            "inicien sesion mediante un formulario con correo electronico y contrasena.",
            "- Validar formato de correo electronico antes de enviar la solicitud.",
            "- En caso de credenciales incorrectas mostrar mensaje de error especifico.",
            "- Generar un JWT firmado con expiracion de 24 horas al autenticar correctamente.",
            "- Almacenar el token en localStorage para persistencia de sesion.",
            "- Redirigir al administrador al panel admin y al cliente a la tienda tras el login exitoso.",
        ]
    )

    req_block(doc,
        "RF-02", "Registro de usuarios",
        "Alta / Esencial",
        [
            "El sistema debe permitir registrar nuevos usuarios con: nombre completo, "
            "correo electronico, contrasena y numero de telefono.",
            "- Validar que el correo no este previamente registrado en el sistema.",
            "- Almacenar la contrasena de forma segura mediante hash bcrypt.",
            "- Asignar por defecto el rol 'cliente' al registrarse.",
            "- Mostrar mensajes de error descriptivos por campo si la validacion falla.",
            "- Los usuarios invitados no requieren registro para navegar el catalogo.",
        ]
    )

    req_block(doc,
        "RF-03", "Control de acceso por roles",
        "Alta / Esencial",
        [
            "El sistema debe implementar control de acceso segun el tipo de usuario:",
            "- Administrador: acceso completo a todos los modulos (productos, pedidos, "
            "usuarios, categorias, ofertas, dashboard y reportes).",
            "- Cliente: acceso a catalogo, carrito, confirmacion de pedido, historial "
            "de ordenes y perfil personal.",
            "- Invitado: acceso de solo lectura al catalogo y ofertas publicas sin "
            "posibilidad de agregar al carrito ni realizar pedidos.",
            "- El control se implementa mediante Guards JWT de NestJS y verificacion "
            "del rol en cada endpoint protegido.",
            "- Las rutas admin del frontend estan protegidas con validacion del token JWT.",
        ]
    )

    req_block(doc,
        "RF-04", "Gestion de productos (CRUD)",
        "Alta / Esencial",
        [
            "El sistema debe permitir al administrador gestionar el inventario de productos:",
            "- Registrar nuevos productos con: nombre, descripcion, precio, cantidad en stock, "
            "categoria, imagen URL y estado (activo/inactivo).",
            "- Editar productos existentes y reflejar cambios en tiempo real en la tienda.",
            "- Eliminar productos, actualizando la disponibilidad en el catalogo.",
            "- Validar que no existan productos duplicados por nombre y referencia.",
            "- Mostrar tabla de productos con busqueda, filtros por categoria y paginacion.",
            "- Alertar cuando el stock de un producto cae por debajo de 5 unidades.",
        ]
    )

    req_block(doc,
        "RF-05", "Catalogo y tienda en linea",
        "Alta / Esencial",
        [
            "El sistema debe mostrar el catalogo de productos de forma publica a todos los usuarios:",
            "- Mostrar productos en grilla de tarjetas con imagen, nombre, precio y categoria.",
            "- Filtrar productos por categoria en tiempo real sin recargar la pagina.",
            "- Busqueda de productos por nombre con resultados instantaneos.",
            "- Mostrar pagina de detalle con imagen ampliada, descripcion completa, "
            "precio y disponibilidad.",
            "- Seccionar los productos en oferta con precios destacados.",
        ]
    )

    req_block(doc,
        "RF-06", "Carrito de compras",
        "Alta / Esencial",
        [
            "El sistema debe permitir que los clientes gestionen su carrito de compras:",
            "- Agregar productos desde el catalogo o la pagina de detalle.",
            "- Visualizar productos en panel lateral (drawer) con imagen, nombre, cantidad y precio.",
            "- Modificar la cantidad de cada producto o eliminarlo del carrito.",
            "- Calcular automaticamente el subtotal y total del carrito.",
            "- Persistir el carrito durante la sesion activa del usuario.",
            "- Vaciar el carrito al confirmar el pedido.",
        ]
    )

    req_block(doc,
        "RF-07", "Confirmacion y gestion de pedidos",
        "Alta / Esencial",
        [
            "El sistema debe permitir que el cliente finalice su compra y que el administrador gestione las ordenes:",
            "- Registrar el pedido con todos los detalles: productos, cantidades, precios y total.",
            "- Generar un numero de orden unico para cada pedido.",
            "- Mostrar pagina de confirmacion con resumen completo.",
            "- Reducir automaticamente el stock de cada producto al confirmar el pedido.",
            "- Permitir al administrador cambiar el estado del pedido "
            "(pendiente, en proceso, enviado, entregado, cancelado).",
            "- Mostrar el historial de pedidos al cliente con estado actualizado.",
        ]
    )

    req_block(doc,
        "RF-08", "Panel administrativo (Dashboard)",
        "Alta / Esencial",
        [
            "El sistema debe proporcionar al administrador un dashboard con indicadores clave:",
            "- Numero total de productos registrados y activos.",
            "- Total de pedidos realizados y su distribucion por estado.",
            "- Total de clientes registrados.",
            "- Ventas totales acumuladas.",
            "- Productos con stock bajo (alerta visual destacada).",
            "- Graficas de ventas por periodo.",
            "- Pedidos recientes con acceso rapido a su detalle.",
        ]
    )

    req_block(doc,
        "RF-09", "Gestion de categorias",
        "Media / Deseado",
        [
            "El sistema debe permitir al administrador gestionar las categorias de productos:",
            "- Crear nuevas categorias con nombre e imagen representativa.",
            "- Editar y eliminar categorias existentes.",
            "- Mostrar las categorias en la tienda como filtros de navegacion.",
            "- Asociar productos a una o mas categorias.",
        ]
    )

    req_block(doc,
        "RF-10", "Gestion de ofertas y promociones",
        "Media / Deseado",
        [
            "El sistema debe permitir al administrador crear y gestionar ofertas:",
            "- Crear ofertas con nombre, descripcion, porcentaje de descuento y fecha de vencimiento.",
            "- Asociar productos o categorias a una oferta.",
            "- Mostrar los productos en oferta con el precio original tachado y el precio con descuento.",
            "- Pagina publica de ofertas accesible sin autenticacion.",
            "- El administrador puede activar o desactivar ofertas.",
        ]
    )

    req_block(doc,
        "RF-11", "Gestion de usuarios (Admin)",
        "Media / Deseado",
        [
            "El sistema debe permitir al administrador gestionar los usuarios registrados:",
            "- Ver la lista de todos los usuarios con nombre, correo, rol y fecha de registro.",
            "- Buscar usuarios por nombre o correo.",
            "- Desactivar o eliminar usuarios del sistema.",
            "- Cambiar el rol de un usuario (cliente a administrador y viceversa).",
        ]
    )

    add_page_break(doc)

    # 3.3 Requisitos no funcionales
    heading2(doc, "3.3 Requisitos No Funcionales")

    # 3.3.1 Rendimiento
    heading3(doc, "3.3.1 Requisitos de Rendimiento")
    body(doc,
        "El sistema debe cumplir con los siguientes criterios de rendimiento medibles:")
    rendimiento = [
        "El 90% de las peticiones a la API REST deben resolverse en menos de 500 ms bajo condiciones normales de uso.",
        "La carga inicial de la aplicacion frontend (SPA) no debe superar 3 segundos en una conexion de 10 Mbps.",
        "El sistema debe soportar al menos 50 usuarios concurrentes sin degradacion apreciable del rendimiento.",
        "Las consultas a la base de datos con filtros y paginacion deben ejecutarse en menos de 200 ms para tablas de hasta 10.000 productos.",
        "Las imagenes de productos deben mostrarse de forma optimizada (lazy loading) para no bloquear la carga inicial.",
        "El dashboard administrativo debe actualizar sus metricas con datos de menos de 1 minuto de antiguedad.",
    ]
    for r in rendimiento:
        bullet(doc, r)

    # 3.3.2 Seguridad
    heading3(doc, "3.3.2 Seguridad")
    body(doc,
        "El sistema implementa las siguientes medidas de seguridad:")
    seguridad_items = [
        "Autenticacion mediante JWT (JSON Web Tokens) con firma HMAC-SHA256 y expiracion configurable.",
        "Las contrasenas se almacenan exclusivamente como hash bcrypt (factor de coste 10+) — nunca en texto plano.",
        "Todos los endpoints protegidos validan el token JWT mediante Guards de NestJS antes de procesar la solicitud.",
        "Las rutas del panel de administracion verifican adicionalmente que el rol del usuario sea 'admin'.",
        "Proteccion contra inyeccion SQL gracias al uso de Prisma ORM con consultas parametrizadas.",
        "Validacion y sanitizacion de todas las entradas del usuario en el backend mediante class-validator.",
        "Configuracion de CORS restrictiva: solo el dominio del frontend en produccion puede consumir la API.",
        "Variables sensibles (JWT_SECRET, DATABASE_URL) gestionadas exclusivamente como variables de entorno — nunca en el codigo fuente.",
        "HTTPS obligatorio en produccion a traves de Render.com (SSL/TLS automatico).",
    ]
    for s in seguridad_items:
        bullet(doc, s)

    # 3.3.3 Fiabilidad
    heading3(doc, "3.3.3 Fiabilidad")
    body(doc,
        "El sistema debe cumplir los siguientes criterios de fiabilidad:")
    fiabilidad = [
        "Disponibilidad objetivo del 99% mensual (less than 7.2 horas de inactividad al mes), garanti­zada por la infraestructura de Render.com.",
        "El sistema debe manejar errores en la API devolviendo codigos HTTP apropiados (400, 401, 403, 404, 500) con mensajes descriptivos.",
        "Los errores internos del servidor deben registrarse en logs para facilitar su diagnostico sin exponer informacion sensible al cliente.",
        "En caso de falla temporal de la base de datos, el sistema debe mostrar un mensaje de error amigable y no colapsar.",
        "El carrito de compras debe persistir su estado durante toda la sesion activa del usuario.",
    ]
    for f in fiabilidad:
        bullet(doc, f)

    # 3.3.4 Disponibilidad
    heading3(doc, "3.3.4 Disponibilidad")
    dispon = [
        "El sistema debe estar disponible el 99% del tiempo en horario de operacion (24/7), sin ventanas de mantenimiento programadas frecuentes.",
        "El despliegue en Render.com garantiza reinicios automaticos del servicio ante caidas.",
        "Las actualizaciones del sistema se realizaran mediante despliegue continuo (CD) desde el repositorio, con tiempo de inactividad minimo.",
        "La base de datos PostgreSQL en Render cuenta con backups automaticos diarios.",
    ]
    for d in dispon:
        bullet(doc, d)

    # 3.3.5 Mantenibilidad
    heading3(doc, "3.3.5 Mantenibilidad")
    mantenib = [
        "El codigo fuente sigue estandares de TypeScript strict con separacion clara en capas (presentacion, negocio, datos).",
        "El backend usa arquitectura modular de NestJS: cada modulo (productos, pedidos, usuarios, auth) es independiente.",
        "El esquema de base de datos se gestiona con migraciones Prisma, lo que permite trazabilidad de cambios.",
        "El frontend usa componentes React atomicos y reutilizables, facilitando modificaciones localizadas.",
        "La documentacion del codigo (README, ARCHITECTURE.md) se mantiene actualizada junto con el codigo fuente.",
        "Las variables de configuracion estan centralizadas en archivos .env para facilitar cambios de entorno.",
    ]
    for m in mantenib:
        bullet(doc, m)

    # 3.3.6 Portabilidad
    heading3(doc, "3.3.6 Portabilidad")
    portab = [
        "El frontend es accesible desde cualquier navegador moderno (Chrome, Firefox, Edge, Safari) sin instalacion adicional.",
        "El backend esta contenerizado con Docker, lo que permite desplegarlo en cualquier plataforma compatible con Docker (AWS, GCP, Azure, VPS).",
        "El sistema no depende de software proprietario: todas las tecnologias usadas (React, NestJS, PostgreSQL, Prisma) son de codigo abierto.",
        "La base de datos puede migrarse a cualquier instancia PostgreSQL compatible simplemente actualizando la variable DATABASE_URL.",
        "El diseno responsivo garantiza compatibilidad con pantallas desde 320px (movil) hasta 2560px (4K).",
        "Mas del 95% del codigo fuente es independiente del sistema operativo del servidor.",
    ]
    for p_item in portab:
        bullet(doc, p_item)

    # 3.4 Otros requisitos
    heading2(doc, "3.4 Otros Requisitos")
    body(doc,
        "Ademas de los requisitos funcionales y no funcionales anteriores, se consideran "
        "los siguientes aspectos complementarios:")

    heading3(doc, "Internacionalizacion")
    bullet(doc, "El sistema opera en idioma espanol (Colombia) como idioma principal.")
    bullet(doc, "Los precios se muestran en pesos colombianos (COP).")
    bullet(doc, "Los formatos de fecha usan el estandar DD/MM/AAAA.")

    heading3(doc, "Accesibilidad")
    bullet(doc, "La interfaz sigue practicas basicas de accesibilidad web (textos alternativos en imagenes, contraste de colores adecuado).")
    bullet(doc, "Los formularios incluyen etiquetas descriptivas (labels) para compatibilidad con lectores de pantalla.")

    heading3(doc, "Privacidad de Datos")
    bullet(doc, "El sistema recopila unicamente los datos necesarios: nombre, correo, telefono y historial de pedidos.")
    bullet(doc, "Los datos personales de los clientes no se comparten con terceros.")
    bullet(doc, "Los usuarios pueden solicitar la eliminacion de su cuenta y datos asociados.")

    add_page_break(doc)

    # ── SECCIÓN 4: APÉNDICES ──────────────────────────────────────────────────
    heading1(doc, "4. Apendices")

    heading2(doc, "Apendice A: Arquitectura de Tres Capas")
    body(doc,
        "StoreHub implementa una arquitectura de tres capas con separacion estricta "
        "de responsabilidades:")
    capas_info = [
        ("Capa de Presentacion (Frontend)",
         "React 19 + TypeScript (SPA)\nVite como bundler\nCSS modular + sistema de diseno propio\nReact Router para navegacion\nContext API para estado global del carrito\nAccesible en: https://inventory-2-sewi.onrender.com"),
        ("Capa de Negocio (Backend)",
         "NestJS + TypeScript\nArquitectura modular (auth, products, orders, users, categories, offers)\nJWT para autenticacion\nclass-validator para validacion de datos\nPrisma ORM para acceso a datos\nAccesible en: https://inventory-1-jkh2.onrender.com/api/v1"),
        ("Capa de Datos (Base de Datos)",
         "PostgreSQL 14+\nPrisma como ORM y gestor de migraciones\nModelos: User, Product, Category, Order, OrderItem, Offer\nAlojada como Managed Database en Render.com"),
    ]
    for capa, desc in capas_info:
        tbl_a = doc.add_table(rows=2, cols=1)
        tbl_a.style = 'Table Grid'
        set_cell_borders(tbl_a)
        set_cell_bg(tbl_a.cell(0, 0), '1A56AB')
        tbl_a.cell(0, 0).paragraphs[0].clear()
        r_capa = tbl_a.cell(0, 0).paragraphs[0].add_run(capa)
        r_capa.bold = True; r_capa.font.size = Pt(10); r_capa.font.color.rgb = BLANCO
        tbl_a.cell(1, 0).paragraphs[0].clear()
        tbl_a.cell(1, 0).paragraphs[0].add_run(desc).font.size = Pt(9)
        doc.add_paragraph()

    heading2(doc, "Apendice B: URLs de Produccion")
    urls = [
        ("Frontend (Tienda y Panel Admin)", "https://inventory-2-sewi.onrender.com"),
        ("Backend API REST",                "https://inventory-1-jkh2.onrender.com/api/v1"),
        ("Repositorio del proyecto",        "GitHub — Repositorio privado del equipo"),
    ]
    for nombre_url, url in urls:
        bold_label(doc, nombre_url, url)

    heading2(doc, "Apendice C: Modelos de Datos Principales")
    body(doc, "Los modelos de datos principales gestionados por Prisma ORM son:")
    modelos = [
        ("User",       "id, nombre, email, password (hash), telefono, role (admin|cliente), createdAt"),
        ("Product",    "id, nombre, descripcion, precio, stock, imagenUrl, estado, categoryId, createdAt"),
        ("Category",   "id, nombre, imagenUrl, createdAt"),
        ("Order",      "id, userId, total, estado, createdAt, updatedAt"),
        ("OrderItem",  "id, orderId, productId, cantidad, precioUnitario"),
        ("Offer",      "id, nombre, descripcion, descuentoPct, activa, fechaFin, productos[]"),
    ]
    tbl_mod = doc.add_table(rows=len(modelos)+1, cols=2)
    tbl_mod.style = 'Table Grid'
    set_cell_borders(tbl_mod)
    table_header_row(tbl_mod, ["Modelo", "Campos principales"])
    for i, (mod, campos) in enumerate(modelos):
        row = tbl_mod.rows[i+1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, 'F0F5FF')
        row.cells[0].paragraphs[0].clear()
        r_mod = row.cells[0].paragraphs[0].add_run(mod)
        r_mod.bold = True; r_mod.font.size = Pt(9)
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run(campos).font.size = Pt(9)

    # Cierre
    doc.add_paragraph()
    p_fin = doc.add_paragraph()
    p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_fin = p_fin.add_run("— Fin del Documento ERS StoreHub v1.0 — Marzo 2026 —")
    r_fin.italic = True
    r_fin.font.size = Pt(9)
    r_fin.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


if __name__ == '__main__':
    print("\n Generando IEEE 830 - ERS StoreHub...\n")
    doc = build_document()
    output = "ERS_StoreHub_IEEE830.docx"
    doc.save(output)
    print(f" Documento generado exitosamente: {output}")
    print(f" Secciones incluidas:")
    print(f"   1. Portada + Ficha del Documento")
    print(f"   2. Introduccion (proposito, alcance, personal, referencias)")
    print(f"   3. Descripcion general (perspectiva, funcionalidad, usuarios, restricciones)")
    print(f"   4. Requisitos especificos (11 RF + 6 RNF)")
    print(f"   5. Apendices (arquitectura, URLs, modelos de datos)")
